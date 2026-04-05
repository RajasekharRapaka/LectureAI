# ═══════════════════════════════════════════════════════════════════════
# LectureAI — app.py
#
# Flask web application that:
#   1. Fetches YouTube transcripts (via yt-dlp, youtube-transcript-api,
#      or direct page scrape — three fallback methods)
#   2. Sends transcripts to Gemini 2.5 Pro/Flash for note generation
#   3. Handles very long transcripts by splitting into chunks and merging
#   4. Runs long jobs in background threads so the browser doesn't timeout
#   5. Exports content as Markdown, HTML, PDF, or DOCX
#   6. Provides a coverage audit endpoint that verifies nothing was missed
# ═══════════════════════════════════════════════════════════════════════

import os
import re
import io
import json
import html
import logging
import math
import subprocess
import tempfile
import threading
import time
import uuid
import xml.etree.ElementTree as ET
from concurrent.futures import ThreadPoolExecutor, as_completed
from pathlib import Path

import requests
from dotenv import load_dotenv
from flask import Flask, render_template, request, jsonify, send_file

# ── Optional export libraries ────────────────────────────────────────
# Imported at module level so helpers like _add_inline_formatting
# can use Pt/RGBColor/Inches without NameError.
# If not installed, the export route returns a clear error message.
try:
    from docx import Document as DocxDocument
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    # Define stubs so the module imports without crashing
    Pt = RGBColor = Inches = WD_ALIGN_PARAGRAPH = None

try:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib import colors as rl_colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Preformatted,
        HRFlowable, ListFlowable, ListItem
    )
    from reportlab.lib.enums import TA_LEFT, TA_CENTER
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

# Load environment variables from the .env file before anything else
load_dotenv()

# ─────────────────────────────────────────────────────────────────────
# LOGGING SETUP
# Writes timestamped log lines to stdout so you can see what the app
# is doing in real time (visible with `python app.py` or `journalctl`).
# ─────────────────────────────────────────────────────────────────────
logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s  %(levelname)-8s  %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
)
logger = logging.getLogger(__name__)

# ─────────────────────────────────────────────────────────────────────
# FLASK APPLICATION
# ─────────────────────────────────────────────────────────────────────
app = Flask(__name__)
app.config["JSON_SORT_KEYS"] = False  # keep JSON keys in insertion order

# ─────────────────────────────────────────────────────────────────────
# CONFIGURATION
# All values are read from environment variables (set in your .env file).
# ─────────────────────────────────────────────────────────────────────

# Your Gemini API key from https://aistudio.google.com/apikey
# Users can also supply their own key per-request via the UI
GEMINI_API_KEY = os.environ.get("GEMINI_API_KEY", "")

# Which Gemini model to use as default (can be overridden per-request)
# Available models (April 2026):
#   gemini-2.5-pro | gemini-2.5-flash | gemini-2.5-flash-lite
#   NOTE: gemini-2.0-flash and gemini-2.0-flash-lite are DEPRECATED & retired March 2026
GEMINI_MODEL = os.environ.get("GEMINI_MODEL", "gemini-2.5-flash")

# Base URL for the Gemini REST API
GEMINI_BASE = "https://generativelanguage.googleapis.com/v1beta/models"

# ── Groq API (free alternative to Gemini — much faster, very generous limits) ─
# Get a free key at: https://console.groq.com
# Free tier: ~14,400 requests/day on llama-3.3-70b (as of April 2026)
GROQ_API_KEY = os.environ.get("GROQ_API_KEY", "")
GROQ_BASE    = "https://api.groq.com/openai/v1/chat/completions"
GROQ_TIMEOUT = 120  # Groq is fast — 120s is more than enough

# Maps UI model names → actual Groq model IDs
GROQ_MODELS = {
    "groq-llama-3.3-70b": "llama-3.3-70b-versatile",   # best quality, ~14k req/day free
    "groq-llama-3.1-8b":  "llama-3.1-8b-instant",       # fastest, very generous limits
    "groq-mixtral":       "mixtral-8x7b-32768",          # good for long context
}

# Path to your exported browser cookies file (Netscape format).
# Used to bypass YouTube's bot detection. Set in .env file.
COOKIES_FILE = os.environ.get("YOUTUBE_COOKIES_FILE", "")

# How many seconds to wait for a YouTube page/caption response
YT_TIMEOUT = 25

# How many seconds to wait for a single Gemini API response.
# 2.5 Pro on a long chunk can take 2-3 minutes.
GEMINI_TIMEOUT = 300

# ── Gemini token / word budgets ──────────────────────────────────────
# Free tier daily limits (April 2026):
#   gemini-2.5-pro:        5 RPM, 100 RPD
#   gemini-2.5-flash:     10 RPM, 250 RPD
#   gemini-2.5-flash-lite: 15 RPM, 1000 RPD  ← most generous Gemini free tier
#   groq-llama-3.3-70b:  ~14,400 RPD          ← best free option overall
#
# We cap output at 32,768 tokens to save ~50% quota vs the 65k max.
MAX_OUTPUT_TOKENS = 32768        # Still ~24,000 words of output — plenty
SAFE_INPUT_WORDS  = 700_000      # We cap input at 700k words to stay safe

# Chunk size: 20k words (~27k tokens input) leaves room for detailed output.
# For Groq models with 128k context, this is fine. For Gemini 1M context, also fine.
CHUNK_WORDS       = 20_000

# Overlap between chunks: last N words of chunk N repeated at start of chunk N+1
CHUNK_OVERLAP_WORDS = 200

# If transcript is longer than this, use async background job
ASYNC_THRESHOLD_WORDS = 12_000

# ── Retry configuration ──────────────────────────────────────────────
GEMINI_MAX_RETRIES   = 4
GEMINI_BASE_WAIT     = 30


def _resolve_api_key(request_key: str = "") -> str:
    """
    Return the effective API key: use the per-request key if provided,
    otherwise fall back to the server's environment variable.
    This lets users supply their own paid-tier key in the UI.
    """
    k = (request_key or "").strip()
    return k if k else GEMINI_API_KEY


def gemini_endpoint(model: str = None, api_key: str = "") -> str:
    """Build the full Gemini API URL for a given model name and key."""
    m   = model or GEMINI_MODEL
    key = api_key or GEMINI_API_KEY
    return f"{GEMINI_BASE}/{m}:generateContent?key={key}"


# ═══════════════════════════════════════════════════════════════════════
# SECTION 1: GENERAL UTILITY FUNCTIONS
# Small helpers used throughout the codebase.
# ═══════════════════════════════════════════════════════════════════════

def extract_video_id(url: str):
    """
    Extract the 11-character YouTube video ID from any URL format.

    Handles:
      - https://www.youtube.com/watch?v=dQw4w9WgXcQ
      - https://youtu.be/dQw4w9WgXcQ
      - https://www.youtube.com/embed/dQw4w9WgXcQ
      - https://www.youtube.com/shorts/dQw4w9WgXcQ
      - dQw4w9WgXcQ  (raw ID)

    Returns the video ID string, or None if not found.
    """
    patterns = [
        r"(?:youtube\.com/watch\?(?:.*&)?v=)([a-zA-Z0-9_-]{11})",
        r"(?:youtu\.be/)([a-zA-Z0-9_-]{11})",
        r"(?:youtube\.com/(?:embed|shorts|v)/)([a-zA-Z0-9_-]{11})",
        r"^([a-zA-Z0-9_-]{11})$",
    ]
    for pattern in patterns:
        match = re.search(pattern, url.strip())
        if match:
            return match.group(1)
    return None


def clean_yt_text(s: str) -> str:
    """
    Decode unicode escape sequences and backslash escapes that YouTube
    embeds in its JSON page data.

    Example input:  "Machine Learning \\u0026 AI"
    Example output: "Machine Learning & AI"
    """
    # Replace \\uXXXX sequences with the actual unicode character
    s = re.sub(r"\\u([0-9a-fA-F]{4})", lambda m: chr(int(m.group(1), 16)), s)
    # Fix other common backslash escapes
    return s.replace("\\n", " ").replace("\\/", "/").replace('\\"', '"').strip()


def strip_code_fences(text: str) -> str:
    """
    Remove markdown code fences that Gemini sometimes wraps its output in.

    Gemini occasionally returns:
        ```html
        <h1>Title</h1>
        ```
    This function strips those fences, returning just the inner content.
    """
    # Remove opening fence (```html, ```markdown, ```md, ``` etc.)
    text = re.sub(r"^```(?:html|markdown|md|json)?\s*\n?", "", text, flags=re.IGNORECASE)
    # Remove closing fence
    text = re.sub(r"\n?```\s*$", "", text)
    return text.strip()


def deduplicate_lines(lines: list) -> list:
    """
    Remove consecutive identical lines from a list.
    This is needed because YouTube's ASR captions often repeat the same
    line twice as the sliding window moves forward.

    Example: ["Hello", "Hello", "world"] → ["Hello", "world"]
    """
    return [line for i, line in enumerate(lines) if i == 0 or line != lines[i - 1]]


def lines_to_transcript(lines: list) -> str:
    """
    Clean a list of caption lines and join them into a single string.

    Steps:
      1. Strip leading/trailing whitespace from each line
      2. Remove empty lines
      3. Deduplicate consecutive identical lines
      4. Join with spaces
      5. Collapse multiple spaces into one
    """
    cleaned = [line.strip() for line in lines if line.strip()]
    deduped = deduplicate_lines(cleaned)
    return re.sub(r"  +", " ", " ".join(deduped)).strip()


def word_count(text: str) -> int:
    """Count the number of words in a string by splitting on whitespace."""
    return len(text.split()) if text.strip() else 0


def estimate_video_hours(wc: int) -> float:
    """
    Estimate video duration in hours from the transcript word count.
    Average speaking pace is ~130 words per minute.

    Example: 78,000 words → 10.0 hours
    """
    return round(wc / 130 / 60, 1)


def chunk_text(text: str, max_words: int = CHUNK_WORDS) -> list:
    """
    Split a long transcript into overlapping chunks for processing.

    Why we need overlap:
      Without overlap, a concept that spans a chunk boundary gets split.
      The first chunk's notes end mid-topic, and the second chunk's notes
      start mid-topic with no context. This causes the "INCOMPLETE" audit result.

    With CHUNK_OVERLAP_WORDS overlap:
      - Last 500 words of chunk N become the first 500 words of chunk N+1
      - The model sees the topic beginning AND continuation in each chunk
      - Merge step deduplicates the overlap

    Returns a list of text strings.
    """
    words = text.split()
    total = len(words)

    # Short enough for single chunk — return as-is
    if total <= max_words:
        return [text]

    chunks = []
    start  = 0

    while start < total:
        end = min(start + max_words, total)

        # Find sentence boundary to avoid cutting mid-sentence
        if end < total:
            segment    = " ".join(words[start:end])
            last_break = max(
                segment.rfind(". ", max(0, len(segment) - 3000)),
                segment.rfind("? ", max(0, len(segment) - 3000)),
                segment.rfind("! ", max(0, len(segment) - 3000)),
            )
            if last_break > 0:
                words_before = len(segment[:last_break + 1].split())
                end          = start + words_before

        chunk = " ".join(words[start:end])
        chunks.append(chunk)

        # Move start forward by (chunk_size - overlap) so next chunk
        # begins CHUNK_OVERLAP_WORDS before the end of this chunk
        advance = (end - start) - CHUNK_OVERLAP_WORDS
        start   = start + max(1, advance)  # always advance at least 1 word

    logger.info(
        "Transcript split into %d chunks (~%d words each, %d-word overlap)",
        len(chunks), max_words, CHUNK_OVERLAP_WORDS
    )
    return chunks


# ═══════════════════════════════════════════════════════════════════════
# SECTION 2: YOUTUBE TRANSCRIPT FETCHING
#
# Three methods tried in order:
#   Method 1: yt-dlp         — most reliable, handles all bot detection
#   Method 2: youtube-transcript-api — good fallback with cookie support
#   Method 3: Direct scrape  — last resort
# ═══════════════════════════════════════════════════════════════════════

def _parse_vtt_subtitles(vtt: str) -> str:
    """
    Parse WebVTT subtitle format into plain text.

    WebVTT looks like this:
        WEBVTT
        00:00:01.000 --> 00:00:03.000
        Hello, welcome to the course.

        00:00:04.000 --> 00:00:06.500
        Today we'll cover neural networks.

    This function strips all the timestamps, headers, and inline tags,
    returning just the spoken text.
    """
    lines = []
    for line in vtt.splitlines():
        line = line.strip()

        # Skip blank lines
        if not line:
            continue

        # Skip VTT header and metadata blocks
        if any(line.startswith(key) for key in
               ("WEBVTT", "NOTE", "Kind:", "Language:", "X-TIMESTAMP")):
            continue

        # Skip timestamp lines like "00:00:01.000 --> 00:00:03.000"
        if "-->" in line or re.match(r"^\d{2}:\d{2}", line):
            continue

        # Skip cue index numbers (just a plain integer on its own line)
        if re.match(r"^\d+$", line):
            continue

        # Strip VTT inline timing tags like <00:00:01.000> and <c>
        line = re.sub(r"<[^>]+>", "", line).strip()

        if line:
            lines.append(line)

    return lines_to_transcript(lines)


def fetch_via_ytdlp(video_id: str) -> tuple:
    """
    Fetch transcript using yt-dlp — the most reliable method.

    yt-dlp is a mature command-line tool that handles:
      - YouTube's bot detection
      - Cookie consent pages
      - Rate limiting
      - Multiple subtitle formats (VTT, SRT, etc.)

    Steps:
      1. Run yt-dlp with --skip-download (don't download the video)
      2. Ask it to write subtitle files (.vtt format)
      3. Parse the .vtt file into plain text

    Returns: (transcript_text, video_title, channel_name)
    Raises: ImportError if yt-dlp not installed
            ValueError with user-friendly message on other failures
    """
    url = f"https://www.youtube.com/watch?v={video_id}"

    # Use a temporary directory that gets automatically cleaned up
    with tempfile.TemporaryDirectory() as tmp_dir:
        output_template = os.path.join(tmp_dir, "sub")

        # Build the yt-dlp command
        cmd = [
            "yt-dlp",
            "--skip-download",           # Don't download the video file
            "--write-auto-sub",          # Download auto-generated captions
            "--write-sub",               # Download manual captions (preferred)
            "--sub-lang", "en,en-US,en-GB",  # English only
            "--sub-format", "vtt",       # Request VTT format
            "--convert-subs", "vtt",     # Convert any other format to VTT
            "--no-playlist",             # Don't process playlist, just this video
            "--quiet",                   # Suppress progress output
            "--no-warnings",             # Suppress warnings
            "-o", output_template,       # Output filename template
            url,
        ]

        # Add cookies file if configured (helps bypass bot detection)
        if COOKIES_FILE and Path(COOKIES_FILE).exists():
            cmd += ["--cookies", COOKIES_FILE]
            logger.info("[yt-dlp] Using cookies file: %s", COOKIES_FILE)

        logger.info("[yt-dlp] Fetching subtitles for video_id=%s", video_id)

        try:
            result = subprocess.run(
                cmd,
                capture_output=True,  # Capture stdout and stderr
                text=True,            # Return strings instead of bytes
                timeout=120,          # Kill if it takes longer than 2 minutes
            )
        except FileNotFoundError:
            # yt-dlp command not found — not installed
            raise ImportError("yt-dlp is not installed. Run: pip install yt-dlp")
        except subprocess.TimeoutExpired:
            raise ValueError("yt-dlp timed out after 2 minutes.")

        # Check for errors
        if result.returncode != 0:
            stderr = result.stderr.strip()
            # Give specific error messages for common failure cases
            if "No subtitles" in stderr or "no subtitles" in stderr.lower():
                raise ValueError("No subtitles found. This video may not have captions.")
            if "Sign in" in stderr or "bot" in stderr.lower():
                raise ValueError(
                    "YouTube blocked the request (bot detection). "
                    "Add your browser cookies — see .env.example for instructions."
                )
            raise ValueError(f"yt-dlp failed: {stderr[:300]}")

        # Find the downloaded .vtt file(s) in the temp directory
        vtt_files = list(Path(tmp_dir).glob("*.vtt"))
        if not vtt_files:
            raise ValueError("yt-dlp ran successfully but produced no subtitle file.")

        # Prefer manual captions over auto-generated ones
        manual_files = [
            f for f in vtt_files
            if ".en." in f.name and "auto" not in f.name.lower()
        ]
        chosen_file = manual_files[0] if manual_files else vtt_files[0]
        logger.info("[yt-dlp] Parsing subtitle file: %s", chosen_file.name)

        # Read and parse the VTT file
        vtt_content = chosen_file.read_text(encoding="utf-8", errors="replace")
        transcript  = _parse_vtt_subtitles(vtt_content)

        if not transcript:
            raise ValueError("Subtitle file was empty after parsing.")

        # Get video metadata separately (title, channel name)
        title, channel = _fetch_ytdlp_metadata(video_id)

        logger.info("[yt-dlp] Success — %d words", word_count(transcript))
        return transcript, title, channel


def _fetch_ytdlp_metadata(video_id: str) -> tuple:
    """
    Get video title and channel name using yt-dlp's --dump-json flag.
    This is a quick call that doesn't download anything.
    Returns ("YouTube Video", "") on failure — metadata is non-critical.
    """
    try:
        cmd = [
            "yt-dlp", "--dump-json", "--no-playlist", "--quiet",
            f"https://www.youtube.com/watch?v={video_id}",
        ]
        if COOKIES_FILE and Path(COOKIES_FILE).exists():
            cmd += ["--cookies", COOKIES_FILE]

        result = subprocess.run(cmd, capture_output=True, text=True, timeout=20)
        if result.returncode == 0 and result.stdout.strip():
            info = json.loads(result.stdout)
            return info.get("title", "YouTube Video"), info.get("uploader", "")
    except Exception as e:
        logger.debug("Metadata fetch failed (non-critical): %s", e)

    return "YouTube Video", ""


def fetch_via_library(video_id: str) -> str:
    """
    Fetch transcript using the youtube-transcript-api Python library.

    This library maintains its own implementation of YouTube's caption
    API and is updated frequently to handle bot-detection changes.
    Supports passing cookies for improved reliability.

    Returns: transcript text string
    Raises: ImportError if library not installed
            ValueError with user-friendly message on failure
    """
    try:
        from youtube_transcript_api import (
            YouTubeTranscriptApi,
            NoTranscriptFound,
            TranscriptsDisabled,
        )
    except ImportError:
        raise ImportError(
            "youtube-transcript-api not installed. Run: pip install youtube-transcript-api"
        )

    # Build keyword arguments — add cookies if available
    kwargs = {}
    if COOKIES_FILE and Path(COOKIES_FILE).exists():
        kwargs["cookies"] = COOKIES_FILE
        logger.info("[library] Using cookies for authentication")

    try:
        api = YouTubeTranscriptApi()

        # First try to get English transcript specifically
        try:
            snippet_list = api.get_transcript(
                video_id,
                languages=["en", "en-US", "en-GB"],
                **kwargs
            )
        except NoTranscriptFound:
            # Fall back to whatever language is available
            logger.info("[library] No English transcript, trying any language")
            all_transcripts = api.list_transcripts(video_id, **kwargs)
            first_available = next(iter(all_transcripts))
            snippet_list    = first_available.fetch()

        # Each snippet is {"text": "...", "start": 1.23, "duration": 2.5}
        lines = [entry.get("text", "").strip() for entry in snippet_list]
        return lines_to_transcript(lines)

    except TranscriptsDisabled:
        raise ValueError("The video owner has disabled transcripts for this video.")
    except NoTranscriptFound:
        raise ValueError("No transcript found — this video may not have captions.")
    except Exception as e:
        raise ValueError(f"youtube-transcript-api error: {e}")


def _make_browser_session() -> requests.Session:
    """
    Create a requests.Session that mimics a real Chrome browser on macOS.

    This is important because:
      - YouTube checks the User-Agent header to detect bots
      - Using a requests.Session (vs bare requests.get) preserves cookies
        across multiple requests — YouTube sets consent cookies on the
        first page load that must be sent with the caption request

    Also loads cookies from the configured cookies file if available.
    """
    session = requests.Session()

    # Set headers that match a real Chrome browser
    session.headers.update({
        "User-Agent": (
            "Mozilla/5.0 (Macintosh; Intel Mac OS X 10_15_7) "
            "AppleWebKit/537.36 (KHTML, like Gecko) "
            "Chrome/124.0.0.0 Safari/537.36"
        ),
        "Accept":          "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
        "Accept-Language": "en-US,en;q=0.9",
        "Accept-Encoding": "gzip, deflate, br",
        "Cache-Control":   "max-age=0",
    })

    # Pre-set consent cookies to skip YouTube's GDPR consent page
    # (without these, YouTube may redirect to a consent form)
    session.cookies.set("CONSENT", "YES+cb.20210328-17-p0.en+FX+667", domain=".youtube.com")
    session.cookies.set("SOCS",    "CAESEwgDEgk2NjY2NjY2NjYaAmVuIAEaBgiA_LysBg",  domain=".youtube.com")

    # Load real browser cookies from file if configured
    if COOKIES_FILE and Path(COOKIES_FILE).exists():
        _load_netscape_cookies(session, COOKIES_FILE)
        logger.info("[scrape] Loaded cookies from %s", COOKIES_FILE)

    return session


def _load_netscape_cookies(session: requests.Session, filepath: str):
    """
    Parse a Netscape-format cookies.txt file and inject all cookies
    into a requests.Session.

    Netscape cookie format (one cookie per line, tab-separated):
        .youtube.com  TRUE  /  FALSE  0  VISITOR_INFO1_LIVE  abc123xyz

    Fields: domain, include_subdomains, path, secure, expiry, name, value
    """
    try:
        with open(filepath, "r", encoding="utf-8", errors="replace") as f:
            for line in f:
                line = line.strip()
                # Skip comments and blank lines
                if not line or line.startswith("#"):
                    continue
                parts = line.split("\t")
                if len(parts) >= 7:
                    domain, _, path, _, _, name, value = parts[:7]
                    session.cookies.set(name, value, domain=domain, path=path)
    except Exception as e:
        logger.warning("Failed to load cookies file (%s): %s", filepath, e)


def _parse_caption_response(body: str) -> str:
    """
    Parse a YouTube caption API response into plain text.

    YouTube can return captions in two formats:
      1. JSON3 format (starts with '{') — newer format
      2. XML/srv3 format (starts with '<') — older format

    Both are handled here.
    """
    body = body.strip()

    # ── JSON3 format ─────────────────────────────────────────────
    # Structure: {"events": [{"segs": [{"utf8": "text"}, ...]}, ...]}
    if body.startswith("{"):
        try:
            data  = json.loads(body)
            lines = []
            for event in data.get("events", []):
                # Each event has segments; concatenate their text
                text = "".join(
                    seg.get("utf8", "") for seg in event.get("segs", [])
                ).strip()
                if text and text != "\n":
                    lines.append(text)
            return lines_to_transcript(lines)
        except Exception as e:
            logger.debug("JSON3 parse failed, trying XML: %s", e)

    # ── XML/srv3 format ───────────────────────────────────────────
    # Structure: <transcript><text start="0" dur="1">Hello</text></transcript>
    try:
        root = ET.fromstring(body)
    except ET.ParseError:
        # Sometimes there's a BOM or garbage before the XML declaration
        cleaned = re.sub(r"^[^<]+", "", body)
        try:
            root = ET.fromstring(cleaned)
        except ET.ParseError as e:
            raise ValueError(f"Cannot parse caption response: {e}")

    lines = []
    for node in root.iter("text"):
        # Get text content and decode HTML entities
        raw = (node.text or "")
        raw = html.unescape(raw)  # decode &amp; &lt; &gt; etc.
        raw = re.sub(r"<[^>]+>", "", raw).strip()  # strip any inline tags
        if raw:
            lines.append(raw)

    return lines_to_transcript(lines)


def fetch_via_scrape(video_id: str) -> tuple:
    """
    Fetch transcript by scraping the YouTube watch page directly.

    This is the fallback method when yt-dlp and the library both fail.
    It works by:
      1. Loading the YouTube watch page (HTML)
      2. Extracting the captionTracks JSON embedded in the page source
      3. Fetching the caption file URL from the track data
      4. Parsing the caption file (XML or JSON3)

    Returns: (transcript_text, video_title, channel_name)
    Raises: ValueError with user-friendly message on failure
    """
    session  = _make_browser_session()
    page_url = f"https://www.youtube.com/watch?v={video_id}"

    # Step 1: Load the YouTube watch page
    logger.info("[scrape] Fetching page for video_id=%s", video_id)
    try:
        response = session.get(page_url, timeout=YT_TIMEOUT)
        response.raise_for_status()
    except requests.exceptions.RequestException as e:
        raise ValueError(f"Cannot reach YouTube: {e}")

    page_html = response.text
    logger.info("[scrape] Page loaded — %d characters", len(page_html))

    # Step 2: Extract video metadata
    title_match   = re.search(r'"title"\s*:\s*"([^"\\]*(?:\\.[^"\\]*)*)"', page_html)
    channel_match = re.search(r'"ownerChannelName"\s*:\s*"([^"\\]*(?:\\.[^"\\]*)*)"', page_html)
    title   = clean_yt_text(title_match.group(1))   if title_match   else "YouTube Video"
    channel = clean_yt_text(channel_match.group(1)) if channel_match else ""

    # Step 3: Extract caption tracks from the embedded JSON
    # YouTube embeds caption track data in the page as a JS object
    caption_tracks = []
    for pattern in [
        r'"captionTracks"\s*:\s*(\[.*?\])',   # standard format
        r'captionTracks\\?":\s*(\[.*?\])',    # escaped variant
    ]:
        match = re.search(pattern, page_html, re.DOTALL)
        if not match:
            continue
        try:
            # Walk the string to find the complete JSON array
            # (can't just use .group(1) because the regex is greedy)
            depth, end = 0, 0
            for i, char in enumerate(match.group(1)):
                if char == "[":
                    depth += 1
                elif char == "]":
                    depth -= 1
                    if depth == 0:
                        end = i + 1
                        break
            caption_tracks = json.loads(match.group(1)[:end])
            logger.info("[scrape] Found %d caption tracks", len(caption_tracks))
            break
        except Exception as e:
            logger.debug("[scrape] Pattern failed: %s", e)
            continue

    # If no tracks found, give a helpful error message
    if not caption_tracks:
        if "consent.youtube.com" in page_html or "CONSENT" in page_html[:3000]:
            raise ValueError(
                "YouTube is showing a consent/bot-check page. "
                "Add your browser cookies to bypass this — see .env.example."
            )
        raise ValueError(
            "No caption tracks found on this video's page. "
            "The video may not have subtitles enabled."
        )

    # Step 4: Pick the best available track
    # Priority: manual English > auto-generated English > any English > anything
    def track_priority(track):
        lang = track.get("languageCode", "")
        kind = track.get("kind", "")
        if lang == "en" and kind != "asr":  return 0  # manual English (best)
        if lang == "en":                    return 1  # auto-generated English
        if lang.startswith("en"):           return 2  # en-GB, en-AU, etc.
        return 3                                      # any other language

    best_track  = sorted(caption_tracks, key=track_priority)[0]
    caption_url = best_track.get("baseUrl", "")

    if not caption_url:
        raise ValueError("Caption track found but has no URL.")

    logger.info("[scrape] Selected track: lang=%s kind=%s",
                best_track.get("languageCode"), best_track.get("kind", "manual"))

    # Step 5: Try fetching captions in both available formats
    # srv3 = XML format, json3 = JSON format
    for fmt in ["srv3", "json3"]:
        # Build URL with the desired format parameter
        url = re.sub(r"[&?]fmt=[^&]*", "", caption_url)  # remove existing fmt param
        url += ("&" if "?" in url else "?") + f"fmt={fmt}"

        try:
            cap_response = session.get(url, timeout=YT_TIMEOUT)
            cap_response.raise_for_status()
            body = cap_response.text

            logger.info("[scrape] Caption response: status=%d len=%d fmt=%s",
                        cap_response.status_code, len(body), fmt)

            if body.strip():
                transcript = _parse_caption_response(body)
                if transcript:
                    return transcript, title, channel

        except Exception as e:
            logger.warning("[scrape] Format %s failed: %s", fmt, e)
            continue

    raise ValueError(
        "YouTube returned empty caption data. "
        "Add your browser cookies — see .env.example for instructions."
    )


def fetch_transcript(video_id: str) -> dict:
    """
    Master transcript fetching function — tries all methods in order.

    Method 1: yt-dlp (most reliable)
    Method 2: youtube-transcript-api (good fallback)
    Method 3: Direct page scrape (last resort)

    Stops at the first success. If all fail, raises ValueError with
    a combined error message explaining what went wrong.

    Returns a dict with:
        transcript      — the full transcript text
        title           — video title
        channel         — channel/uploader name
        word_count      — number of words in transcript
        estimated_hours — estimated video length in hours
        chunks_needed   — how many processing chunks this will need
        method          — which method succeeded
    """
    errors  = []  # collect errors from each method for the final error message
    title   = "YouTube Video"
    channel = ""

    # ── Method 1: yt-dlp ─────────────────────────────────────────
    try:
        logger.info("[fetch] Trying yt-dlp for %s", video_id)
        transcript, title, channel = fetch_via_ytdlp(video_id)
        wc = word_count(transcript)
        logger.info("[fetch] yt-dlp succeeded — %d words", wc)
        return _build_transcript_result(transcript, title, channel, wc, "yt-dlp")
    except ImportError:
        errors.append("yt-dlp: not installed")
        logger.info("[fetch] yt-dlp not installed, trying next method")
    except ValueError as e:
        errors.append(f"yt-dlp: {e}")
        logger.warning("[fetch] yt-dlp failed: %s", e)

    # ── Method 2: youtube-transcript-api ─────────────────────────
    try:
        logger.info("[fetch] Trying youtube-transcript-api for %s", video_id)
        transcript = fetch_via_library(video_id)

        # Library doesn't give us metadata — try to fetch it separately
        try:
            meta_session  = _make_browser_session()
            meta_response = meta_session.get(
                f"https://www.youtube.com/watch?v={video_id}", timeout=YT_TIMEOUT
            )
            if meta_response.ok:
                page = meta_response.text
                tm   = re.search(r'"title"\s*:\s*"([^"\\]*(?:\\.[^"\\]*)*)"', page)
                cm   = re.search(r'"ownerChannelName"\s*:\s*"([^"\\]*(?:\\.[^"\\]*)*)"', page)
                if tm: title   = clean_yt_text(tm.group(1))
                if cm: channel = clean_yt_text(cm.group(1))
        except Exception:
            pass  # metadata failure is non-critical

        wc = word_count(transcript)
        logger.info("[fetch] youtube-transcript-api succeeded — %d words", wc)
        return _build_transcript_result(transcript, title, channel, wc, "youtube-transcript-api")
    except ImportError:
        errors.append("youtube-transcript-api: not installed")
        logger.info("[fetch] Library not installed, trying scrape")
    except ValueError as e:
        errors.append(f"youtube-transcript-api: {e}")
        logger.warning("[fetch] Library failed: %s", e)

    # ── Method 3: Direct scrape ───────────────────────────────────
    try:
        logger.info("[fetch] Trying direct page scrape for %s", video_id)
        transcript, title, channel = fetch_via_scrape(video_id)
        wc = word_count(transcript)
        logger.info("[fetch] Scrape succeeded — %d words", wc)
        return _build_transcript_result(transcript, title, channel, wc, "scrape")
    except ValueError as e:
        errors.append(f"scrape: {e}")
        logger.warning("[fetch] Scrape failed: %s", e)

    # All methods failed — raise with combined error details
    raise ValueError(
        "Could not fetch transcript after trying all methods. "
        "The most common fix is adding your browser cookies — see .env.example.\n"
        "Details: " + " | ".join(errors)
    )


def _build_transcript_result(transcript, title, channel, wc, method) -> dict:
    """Helper to build the standardized transcript response dict."""
    return {
        "transcript":       transcript,
        "title":            title,
        "channel":          channel,
        "word_count":       wc,
        "estimated_hours":  estimate_video_hours(wc),
        "chunks_needed":    math.ceil(wc / CHUNK_WORDS),
        "method":           method,
    }


# ═══════════════════════════════════════════════════════════════════════
# SECTION 3: AI API — GEMINI + GROQ
# call_gemini() is the universal entry point — it routes to Groq or Gemini
# automatically based on the model name prefix.
# ═══════════════════════════════════════════════════════════════════════

def call_groq(
    system_instruction: str,
    user_prompt:        str,
    model:              str,
    api_key:            str = "",
) -> str:
    """
    Call the Groq API (OpenAI-compatible format) and return the text response.

    Groq is a free alternative to Gemini — much faster (LPU hardware) and
    extremely generous free tier (~14,400 requests/day on llama-3.3-70b).

    Free tier (April 2026):
      llama-3.3-70b-versatile: ~14,400 req/day, best quality
      llama-3.1-8b-instant:    very high limits, fastest
      mixtral-8x7b-32768:      ~14,400 req/day, 32k context

    Get a free key at: https://console.groq.com
    """
    key = (api_key or "").strip() or GROQ_API_KEY
    if not key:
        raise ValueError(
            "No Groq API key configured.\n"
            "Get a free key at https://console.groq.com\n"
            "Then add GROQ_API_KEY=gsk_... to your .env file,\n"
            "or paste it in the 🔑 API Key settings panel."
        )

    groq_model = GROQ_MODELS.get(model, "llama-3.3-70b-versatile")
    logger.info("[Groq] Calling %s (→ %s)", model, groq_model)

    payload = {
        "model":    groq_model,
        "messages": [
            {"role": "system", "content": system_instruction},
            {"role": "user",   "content": user_prompt},
        ],
        "max_tokens":  8192,
        "temperature": 0.0,
    }

    for attempt in range(GEMINI_MAX_RETRIES):
        try:
            response = requests.post(
                GROQ_BASE,
                headers={
                    "Authorization": f"Bearer {key}",
                    "Content-Type":  "application/json",
                },
                json=payload,
                timeout=GROQ_TIMEOUT,
            )
        except requests.exceptions.Timeout:
            raise TimeoutError(f"Groq API timed out after {GROQ_TIMEOUT}s.")
        except requests.exceptions.RequestException as e:
            raise ConnectionError(f"Network error calling Groq API: {e}")

        if response.ok:
            text = response.json()["choices"][0]["message"]["content"]
            if not text:
                raise RuntimeError("Groq returned an empty response.")
            return text

        err_msg = response.text
        try:
            err_msg = response.json().get("error", {}).get("message", response.text)
        except Exception:
            pass

        if response.status_code in (429, 413) and attempt < GEMINI_MAX_RETRIES - 1:
            retry_after = response.headers.get("Retry-After")
            # 413 = TPM (tokens-per-minute) exceeded — wait 65s for the window to reset
            default_wait = 65 if response.status_code == 413 else GEMINI_BASE_WAIT * (2 ** attempt)
            wait_secs    = int(retry_after) + 3 if retry_after else default_wait
            logger.warning("Groq rate limit (%d) — waiting %ds (attempt %d/%d).",
                           response.status_code, wait_secs, attempt + 1, GEMINI_MAX_RETRIES)
            time.sleep(wait_secs)
            continue

        if response.status_code == 401:
            raise RuntimeError(
                "Groq API key is invalid (401). "
                "Check your key at https://console.groq.com"
            )

        raise RuntimeError(f"Groq API error {response.status_code}: {err_msg[:200]}")

    raise RuntimeError(
        f"Groq rate limit persists after {GEMINI_MAX_RETRIES} retries. "
        "Wait a few minutes or switch to a different model."
    )


# ═══════════════════════════════════════════════════════════════════════

def call_gemini(
    system_instruction: str,
    user_prompt:        str,
    model:              str = None,
    api_key:            str = "",
) -> str:
    """
    Universal AI call — routes to Groq OR Gemini based on model name prefix.

    Groq models (model starts with "groq-"):
      groq-llama-3.3-70b  — best quality, ~14,400 free req/day ⭐
      groq-llama-3.1-8b   — fastest, very high free limits
      groq-mixtral        — good for longer content

    Gemini models (April 2026 free tier):
      gemini-2.5-pro:        100 req/day
      gemini-2.5-flash:      250 req/day
      gemini-2.5-flash-lite: 1,000 req/day

    Smart 429 handling for Gemini:
      - Daily quota exhausted → fail immediately with clear message
      - RPM limit             → wait 30s/60s/120s/240s and retry
    """
    model = model or GEMINI_MODEL

    # ── Route to Groq for groq- prefixed models ────────────────────
    if model.startswith("groq-"):
        return call_groq(system_instruction, user_prompt, model, api_key)

    # ── Gemini path ────────────────────────────────────────────────
    effective_key = _resolve_api_key(api_key)
    if not effective_key:
        raise ValueError(
            "No API key configured. Options:\n"
            "  • Set GEMINI_API_KEY in your .env file (https://aistudio.google.com/apikey)\n"
            "  • Set GROQ_API_KEY in your .env file and use a Groq model (https://console.groq.com)\n"
            "  • Enter your key in the 🔑 API Key settings panel in the app"
        )

    payload = {
        "system_instruction": {
            "parts": [{"text": system_instruction}]
        },
        "contents": [
            {"role": "user", "parts": [{"text": user_prompt}]}
        ],
        "generationConfig": {
            "maxOutputTokens": MAX_OUTPUT_TOKENS,
            "temperature":     0.0,
            "topP":            0.95,
        },
    }

    endpoint        = gemini_endpoint(model, effective_key)
    RETRYABLE_CODES = {429, 500, 503}

    for attempt in range(GEMINI_MAX_RETRIES):
        try:
            response = requests.post(endpoint, json=payload, timeout=GEMINI_TIMEOUT)
        except requests.exceptions.Timeout:
            raise TimeoutError(
                f"Gemini API timed out after {GEMINI_TIMEOUT}s. "
                "Try gemini-2.5-flash-lite or a Groq model (fastest)."
            )
        except requests.exceptions.RequestException as e:
            raise ConnectionError(f"Network error calling Gemini API: {e}")

        # ── Success ────────────────────────────────────────────────
        if response.ok:
            break

        # ── Parse error details ────────────────────────────────────
        err_msg    = response.text
        err_status = ""
        try:
            err_body   = response.json()
            err_msg    = err_body.get("error", {}).get("message", response.text)
            err_status = err_body.get("error", {}).get("status", "")
        except Exception:
            pass

        # ── Smart 429 handling ─────────────────────────────────────
        if response.status_code == 429:
            msg_lower = err_msg.lower()

            # Check the HTTP Retry-After header — if it's short (<120s), it's RPM not daily quota
            retry_after_header = response.headers.get("Retry-After")
            is_short_retry     = retry_after_header and int(retry_after_header) < 120

            # Detect DAILY quota exhaustion — retrying won't help until tomorrow
            is_daily_exhausted = (
                not is_short_retry and
                any(kw in msg_lower for kw in [
                    "per day", "perday", "daily", "per_day",
                    "generatecontentrequest",      # quota metric name in errors
                    "quota exceeded",
                    "exceeded your current quota",
                    "resource_exhausted",          # some 2.0 models return this directly
                    "requests per day",
                    "billing",                     # billing/plan issues
                ])
            )

            # If the 429 body looks like a generic RPM message, treat as retryable
            is_rpm = any(kw in msg_lower for kw in [
                "resource has been exhausted",
                "requests per minute",
                "rate limit",
                "too many requests",
            ])

            if is_rpm:
                is_daily_exhausted = False

            if is_daily_exhausted:
                if "pro" in model:
                    alt_model = "gemini-2.5-flash (250/day) or groq-llama-3.3-70b (~14k/day)"
                elif "flash-lite" in model:
                    alt_model = "groq-llama-3.3-70b (~14,400/day free) — switch to Groq!"
                else:
                    alt_model = "gemini-2.5-flash-lite (1,000/day) or groq-llama-3.3-70b (~14k/day)"

                key_note = "custom key" if api_key else "server env key"
                raise RuntimeError(
                    f"🚫 Daily API quota exhausted for '{model}' ({key_note}).\n\n"
                    "Options:\n"
                    f"  • Switch model → {alt_model}\n"
                    "  • Try Groq — free key at https://console.groq.com (~14k req/day)\n"
                    "  • Wait until tomorrow (Gemini resets ~midnight PT)\n"
                    "  • Add/change your key in the 🔑 API Key settings panel\n\n"
                    f"Error: {err_msg[:120]}"
                )

            # RPM rate limit — worth retrying after a short wait
            if attempt < GEMINI_MAX_RETRIES - 1:
                retry_after = response.headers.get("Retry-After")
                wait_secs   = int(retry_after) + 3 if retry_after else GEMINI_BASE_WAIT * (2 ** attempt)
                logger.warning(
                    "Gemini RPM limit (429) for %s — waiting %ds (attempt %d/%d).",
                    model, wait_secs, attempt + 1, GEMINI_MAX_RETRIES
                )
                time.sleep(wait_secs)
                continue

        # ── Other retryable errors (500, 503) ──────────────────────
        if response.status_code in RETRYABLE_CODES and attempt < GEMINI_MAX_RETRIES - 1:
            wait_secs = 15 * (attempt + 1)
            logger.warning("Gemini %d — waiting %ds (attempt %d/%d).",
                           response.status_code, wait_secs, attempt + 1, GEMINI_MAX_RETRIES)
            time.sleep(wait_secs)
            continue

        # ── Non-retryable / last attempt ───────────────────────────
        if response.status_code == 400:
            raise RuntimeError(
                f"Gemini rejected request (400): {err_msg[:300]}. "
                "Input may be too large — try a shorter transcript or smaller chunk size."
            )
        if response.status_code == 403:
            raise RuntimeError(
                f"Gemini access denied (403). Verify your API key is valid. "
                f"Detail: {err_msg[:200]}"
            )
        raise RuntimeError(f"Gemini API error {response.status_code}: {err_msg[:300]}")

    else:
        raise RuntimeError(
            f"Gemini RPM limit persists after {GEMINI_MAX_RETRIES} retries for '{model}'.\n"
            "Options:\n"
            "  • Wait a few minutes and try again\n"
            "  • Switch to gemini-2.5-flash-lite (15 RPM) or a Groq model\n"
            "  • Try Groq (free at https://console.groq.com) — much higher limits\n"
            "  • Enter your own key in 🔑 Settings for higher limits"
        )

    data = response.json()

    # Check why the model stopped generating
    finish_reason = data.get("candidates", [{}])[0].get("finishReason", "")
    if finish_reason == "SAFETY":
        raise RuntimeError("Gemini blocked this response due to safety filters.")
    if finish_reason == "RECITATION":
        raise RuntimeError("Gemini blocked this response due to recitation policy.")
    if finish_reason == "MAX_TOKENS":
        # Output was truncated — log warning but continue (partial notes are OK)
        logger.warning("Gemini hit MAX_TOKENS limit — output may be truncated at the end")

    # Extract the text from the nested response structure
    text = (
        data.get("candidates", [{}])[0]
            .get("content", {})
            .get("parts", [{}])[0]
            .get("text", "")
    )

    if not text:
        raise RuntimeError(
            "Gemini returned an empty response. "
            "Check your API key is valid and has quota remaining."
        )

    return text


# ═══════════════════════════════════════════════════════════════════════
# SECTION 4: CHUNKED TRANSCRIPT PROCESSING
#
# For long lectures, we:
#   1. Split transcript into ~60k-word chunks
#   2. Generate notes/article for each chunk independently
#   3. Merge all chunk outputs into one coherent document
# ═══════════════════════════════════════════════════════════════════════

# System prompt used for all note-generation calls.
# This is the most important prompt — it determines note quality.
SYSTEM_NOTES = """You are the world's most thorough technical note-taker and educator.
You specialize in data science, machine learning, AI, software engineering, and any technical topic.

═══════════════════════════════════════════════════════════
YOUR SINGLE MOST IMPORTANT RULE — READ THIS FIRST:
═══════════════════════════════════════════════════════════
You must produce EXHAUSTIVE, DETAILED notes — NOT a summary.
A summary condenses. Notes EXPAND and EXPLAIN every single point.
If your output looks like a paragraph summary, you have FAILED.

YOUR OUTPUT MUST:
✅ Cover EVERY concept, fact, term, number, name, and example from the transcript
✅ Define every technical term fully — not just name it
✅ Write out every formula with all variables explained
✅ Include every code snippet with explanation
✅ Reproduce every analogy and example the instructor gives
✅ Capture every specific number, cost, dataset name, algorithm name mentioned
✅ Use ## headings, ### subheadings, bullet points, and bold for structure
✅ Be AT LEAST 2000 words for every 10 minutes of transcript content

YOUR OUTPUT MUST NOT:
❌ Summarize a topic in one sentence when the instructor spent 3 minutes on it
❌ Use vague phrases like "and more", "etc.", "various techniques"
❌ Skip anything because it "seems minor" — everything the instructor says is important
❌ Collapse multiple distinct concepts into one vague phrase
❌ Write fewer than 3000 words for a 30-minute transcript chunk

CRITICAL SPECIFICS TO ALWAYS CAPTURE:
- Every specific number mentioned (costs, parameter counts, token counts, percentages)
- Every named algorithm (SHA-1, MD5, MinHash, BPE, RLHF, DPO — whatever is mentioned)
- Every named dataset, model, project, or company
- Every step-by-step process — numbered
- Every "why" the instructor explains — not just the "what"
- Every comparison ("X vs Y", "X is better than Y because...")
- Every warning, pitfall, or "don't do this" statement

OUTPUT FORMAT:
## [Topic Name]

**[Subtopic]**: Full explanation here. Define it. Explain why it matters. Explain how it works.

- Specific fact or detail with actual names/numbers
- Another specific fact

> ⚠️ **Important**: Any warning or pitfall the instructor mentions

```
Any code or formula block
```

**Key Takeaways**:
- Takeaway 1 with specifics
- Takeaway 2 with specifics"""


def _get_notes_system_prompt(tone: str) -> str:
    """Return the appropriate system prompt for a given note style/tone."""
    if tone in ("concise", "executive"):
        return (
            "You are an expert academic summarizer. Convert lecture transcripts into clear, "
            "accurate, concise summaries. Be brief and focused — only the most essential points. "
            "Never use phrases like 'In this video' or 'The speaker says'."
        )
    elif tone == "bullet":
        return (
            "You are an expert note-taker. Extract the most important facts, concepts, and "
            "takeaways from lecture transcripts as clean bullet points. "
            "Be specific and informative — no vague generalities."
        )
    elif tone == "detailed":
        return (
            "You are an expert academic summarizer. Convert lecture transcripts into dense, "
            "well-organized paragraph summaries covering all major themes and techniques. "
            "Never use phrases like 'In this video' or 'The speaker says'."
        )
    elif tone == "notes":
        return (
            "You are an expert technical note-taker. Create structured lecture notes with "
            "clear headings, key definitions, formulas, and the most important takeaways. "
            "Capture the essential content — organized and scannable."
        )
    else:  # comprehensive (default)
        return SYSTEM_NOTES


def _build_notes_prompt(chunk: str, chunk_num: int, total_chunks: int, title: str = "", tone: str = "comprehensive") -> str:
    """Build the per-chunk notes generation prompt based on selected tone."""
    topic_line  = f"Course: {title}\n" if title else ""
    part_line   = f"TRANSCRIPT PART {chunk_num} OF {total_chunks}\n" if total_chunks > 1 else ""
    overlap_note = (
        "\nNOTE: This chunk may begin mid-topic (it overlaps with the previous chunk). "
        "Start your notes from wherever new content begins.\n"
        if chunk_num > 1 else ""
    )

    # ── Tone-specific output instructions ─────────────────────────────
    if tone == "concise":
        output_instructions = """Write a concise 3-5 sentence summary of the key message and main points.
Focus only on the single most important insight and practical takeaways.
Do NOT use headings, bullet points, or lists. Write flowing prose only.

Begin summary:"""

    elif tone == "executive":
        output_instructions = """Write exactly 3 sentences:
1. What topic/problem this lecture covers
2. The key insight, technique, or solution presented
3. The main actionable conclusion or takeaway

Do NOT add anything else. Three sentences only.

Begin:"""

    elif tone == "bullet":
        output_instructions = """Write 8-15 bullet points covering the most important facts, concepts,
algorithms, formulas, and practical insights from this transcript.

Rules:
- Each bullet must be specific and informative (not vague)
- Include actual terms, names, formulas where relevant
- Order from most to least important
- No sub-bullets — keep it flat and scannable

Format:
• [Key point 1]
• [Key point 2]
...

Begin bullet points:"""

    elif tone == "detailed":
        output_instructions = """Write a detailed 3-4 paragraph summary covering:
- All major themes and arguments presented
- Key techniques, algorithms, or concepts introduced
- Important examples or case studies mentioned
- Main takeaways and conclusions

Write in flowing prose — no bullet points or headers.
Each paragraph should cover a distinct aspect of the lecture.

Begin summary:"""

    elif tone == "notes":
        output_instructions = """Write structured lecture notes with clear headings and sub-points.

STRUCTURE:
## [Main Topic]
### [Subtopic]
- Key point or definition
- Formula or algorithm (if any)
- Important example (if any)

Include:
- All major concepts with brief definitions
- Key formulas written out with variable explanations
- Important code snippets (if mentioned)
- Notable warnings or common mistakes the instructor flags

Keep it organized and scannable — not exhaustive, but complete on essentials.

Begin notes:"""

    else:  # comprehensive (default)
        wc_estimate = len(chunk.split())
        min_words   = max(2000, wc_estimate // 3)  # notes should be at least 1/3 the transcript size
        output_instructions = f"""YOU ARE WRITING DETAILED LECTURE NOTES — NOT A SUMMARY.

MINIMUM LENGTH REQUIREMENT: Your notes MUST be at least {min_words:,} words.
If you finish before reaching this length, you have skipped content. Go back and expand.

═══════════════════════════════════════════════════════════
STEP 1: SCAN THE TRANSCRIPT FOR EVERY SPECIFIC FACT
═══════════════════════════════════════════════════════════
Before writing, mentally list everything mentioned:
- Every named algorithm, technique, or method (e.g., SHA-1, MinHash, BPE, RLHF, DPO)
- Every specific number, cost, size, or metric (e.g., $100M, 1.76T parameters, 70T tokens)
- Every named dataset, model, project, or company (e.g., FineWeb, Alpaca, Stargate)
- Every step-by-step process
- Every analogy or example the instructor gives
- Every comparison or trade-off explained
- Every warning or pitfall mentioned

═══════════════════════════════════════════════════════════
STEP 2: WRITE EXHAUSTIVE NOTES IN THIS EXACT STRUCTURE
═══════════════════════════════════════════════════════════

## [Main Topic — use the actual topic name from the transcript]

### [Subtopic Name]

**What it is**: [Full definition — not just a label. Explain what this actually means.]

**Why it matters**: [Why does the instructor teach this? What problem does it solve?]

**How it works**: [Step-by-step. Number every step. Include every detail mentioned.]
1. Step one — exactly as explained
2. Step two — with any specific values or algorithms named
3. ...

**Specific facts & numbers**:
- [Every specific number, name, cost, size mentioned — e.g., "GPT-4: 1.76 trillion parameters"]
- [Named algorithms with their purpose — e.g., "SHA-1 and MD5: hash functions for exact deduplication"]
- [Named datasets/projects — e.g., "FineWeb: open-source dataset with 25.9 billion rows"]

**Example from lecture**: [Reproduce any example the instructor walks through, fully worked out]

**Instructor's analogy**: [Include word-for-word if given — analogies aid memory]

> ⚠️ **Warning / Common Mistake**: [Any pitfall or "don't do this" the instructor mentions]

---

═══════════════════════════════════════════════════════════
STEP 3: COMPLETENESS AUDIT BEFORE SUBMITTING
═══════════════════════════════════════════════════════════
After writing, re-read the transcript and check:
□ Did I capture every specific number mentioned?
□ Did I define every named algorithm or technique?
□ Did I include every named dataset, model, or project?
□ Did I reproduce every example and analogy?
□ Did I capture every step of every process?
□ Is my output at least {min_words:,} words?

If you answered NO to any of these, add the missing content NOW before submitting.

BEGIN YOUR COMPREHENSIVE NOTES BELOW:
"""

    return f"""{topic_line}{part_line}{overlap_note}
=== TRANSCRIPT ===
{chunk}
=== END TRANSCRIPT ===

{output_instructions}"""


def _build_merge_prompt(chunk_notes: list, title: str, total_chunks: int) -> str:
    """
    Build the merge prompt for combining multiple chunk notes.
    For very large note sets (>80k words), we do a hierarchical merge
    in process_transcript to avoid exceeding input limits.
    """
    total_words = sum(word_count(n) for n in chunk_notes)
    parts_text  = "\n\n════════════════ PART BREAK ════════════════\n\n".join(
        f"# PART {i+1} NOTES\n\n{notes}"
        for i, notes in enumerate(chunk_notes)
    )

    return f"""You have {total_chunks} sets of lecture notes to merge into ONE perfect document.
Course: "{title}"
Total content: ~{total_words:,} words across {total_chunks} parts.

YOUR TASK: Merge ALL parts into ONE comprehensive, perfectly organized study document.

RULES — follow every single one:
1. KEEP EVERYTHING — Do not drop, shorten, or summarize ANY content from any part
2. DEDUPLICATE CAREFULLY — Same concept in multiple parts? Keep the MOST DETAILED version
3. REORGANIZE LOGICALLY — If a topic spans multiple parts, group it together
4. STRUCTURE REQUIRED:
   a) Start: "## Table of Contents" with every section listed
   b) Then: "## Prerequisites" (if lecture implies prior knowledge)
   c) Then: All the lecture content, organized logically
   d) End: "## Master Summary" (5-7 comprehensive paragraphs)
   e) End: "## Quick Reference Cheat Sheet" (key formulas, terms, code)
5. SMOOTH FLOW — Add brief transitions between sections
6. FORMAT: Consistent Markdown with proper ## / ### / #### hierarchy

THE MERGED DOCUMENT MUST BE AT LEAST AS LONG AS THE SUM OF ALL PARTS.
Every formula, every example, every analogy, every warning must survive.

═══════════════ INPUT PARTS ═══════════════

{parts_text}

═══════════════════════════════════════════

Begin the merged document with ## Table of Contents:"""


def _build_article_prompt(
    chunk: str, fmt_desc: str, fmt_instructions: str,
    chunk_num: int, total_chunks: int, title: str = ""
) -> str:
    """
    Build the user prompt for generating an article from a transcript chunk.
    Similar structure to the notes prompt but for article formats.
    """
    part_info = f"(Part {chunk_num} of {total_chunks})" if total_chunks > 1 else ""
    return f"""Course/Lecture: {title} {part_info}

Create {fmt_desc} from the following transcript segment.
This must be comprehensive — capture every concept, technique, and insight.
A reader should be able to fully learn the material without watching the video.

TRANSCRIPT:
{chunk}

FORMAT INSTRUCTIONS:
{fmt_instructions}

Begin now:"""


def process_transcript(
    transcript:   str,
    mode:         str,    # "notes" or "article"
    tone_or_fmt:  str,    # tone name (for notes) or format key (for articles)
    title:        str = "",
    model:        str = None,
    api_key:      str = "",   # per-request API key (overrides env var)
    progress_cb   = None,  # optional callback(message: str) for progress updates
) -> dict:
    """
    Process a transcript of any length into notes or an article.

    For short transcripts: single API call
    For long transcripts:  multiple chunked calls + merge pass

    Parameters:
        transcript  — the full transcript text
        mode        — "notes" to generate study notes, "article" for formatted articles
        tone_or_fmt — the style/format key (e.g., "comprehensive", "html", "markdown")
        title       — video title for context
        model       — Gemini model to use (defaults to GEMINI_MODEL env var)
        progress_cb — optional function called with status strings during processing

    Returns a dict with:
        content    — the generated text (Markdown or HTML)
        chunks     — how many chunks were processed
        words_in   — input word count
        words_out  — output word count
        model_used — which model was used
    """
    model      = model or GEMINI_MODEL
    wc         = word_count(transcript)
    # Groq free tier has ~12k TPM limit; system prompt + notes prompt overhead
    # uses ~1-2k tokens, so cap transcript chunks at ~7k words (~9k tokens) for Groq.
    chunk_words = 7_000 if model.startswith("groq-") else CHUNK_WORDS
    chunks     = chunk_text(transcript, chunk_words)
    num_chunks = len(chunks)

    def log_progress(msg):
        """Send progress updates to both the logger and the optional callback."""
        logger.info(msg)
        if progress_cb:
            progress_cb(msg)

    log_progress(
        f"Processing: {wc:,} words → {num_chunks} chunk(s), "
        f"model={model}, mode={mode}"
    )

    # ── Process chunks (parallel for Groq, sequential for Gemini) ─
    chunk_outputs = [None] * num_chunks  # pre-allocate to preserve order

    # Groq has high RPM limits → parallel is safe and 3-4× faster
    # Gemini free tier has low RPM (5-15) → sequential avoids 429s
    use_parallel = model.startswith("groq-") and num_chunks > 1
    max_workers  = min(num_chunks, 4) if use_parallel else 1

    def _process_one_chunk(idx_chunk):
        idx, chunk = idx_chunk
        chunk_wc = word_count(chunk)
        log_progress(f"Chunk {idx+1}/{num_chunks}: {chunk_wc:,} words → {model}…")

        if mode == "notes":
            prompt     = _build_notes_prompt(chunk, idx+1, num_chunks, title, tone_or_fmt)
            sys_prompt = _get_notes_system_prompt(tone_or_fmt)
            output     = call_gemini(sys_prompt, prompt, model, api_key)
        elif mode == "article":
            fmt_config = FORMAT_MAP.get(tone_or_fmt, FORMAT_MAP["html"])
            prompt     = _build_article_prompt(
                chunk, fmt_config["desc"], fmt_config["instructions"],
                idx+1, num_chunks, title
            )
            sys_prompt = (
                "You are an expert technical content writer creating educational articles "
                "from lecture transcripts. Be comprehensive — capture everything taught."
            )
            output = strip_code_fences(call_gemini(sys_prompt, prompt, model, api_key))
        else:
            raise ValueError(f"Unknown mode: '{mode}'. Use 'notes' or 'article'.")

        log_progress(f"Chunk {idx+1}/{num_chunks} complete — {word_count(output):,} words generated.")
        return idx, output

    if use_parallel:
        log_progress(f"⚡ Running {num_chunks} chunks in parallel ({max_workers} workers)…")
        with ThreadPoolExecutor(max_workers=max_workers) as executor:
            futures = {executor.submit(_process_one_chunk, (i, c)): i
                       for i, c in enumerate(chunks)}
            for future in as_completed(futures):
                idx, output = future.result()
                chunk_outputs[idx] = output
    else:
        for idx, chunk in enumerate(chunks):
            i, output = _process_one_chunk((idx, chunk))
            chunk_outputs[i] = output

    # ── Merge all chunks into final document ──────────────────────
    if num_chunks == 1:
        # No merging needed — just use the single output
        final_output = chunk_outputs[0]
        log_progress("Single chunk — no merge needed.")
    else:
        log_progress(f"Merging {num_chunks} chunk notes…")

        if mode == "notes":
            is_short_format = tone_or_fmt in ("concise", "executive", "bullet", "detailed")

            if is_short_format:
                if tone_or_fmt == "bullet":
                    all_parts = "\n".join(chunk_outputs)
                    merge_prompt = (
                        f"Combine these bullet point lists from {num_chunks} parts about '{title}' "
                        f"into ONE clean deduplicated list of 10-20 key bullets. "
                        f"Keep only the most specific, informative points.\n\n{all_parts}\n\nFinal bullet list:"
                    )
                    final_output = call_gemini(
                        "You are an expert summarizer. Combine bullet lists precisely.",
                        merge_prompt, model, api_key
                    )
                elif tone_or_fmt in ("concise", "executive"):
                    all_parts = "\n\n---\n\n".join(chunk_outputs)
                    sentences = "3-5 sentences" if tone_or_fmt == "concise" else "exactly 3 sentences"
                    merge_prompt = (
                        f"These are summaries of {num_chunks} parts of a lecture about '{title}':\n\n"
                        f"{all_parts}\n\n"
                        f"Write ONE unified {sentences} summary of the complete lecture. Output ONLY the summary:"
                    )
                    final_output = call_gemini("You are an expert summarizer.", merge_prompt, model, api_key)
                else:  # detailed
                    all_parts = "\n\n---\n\n".join(chunk_outputs)
                    merge_prompt = (
                        f"These are summaries of {num_chunks} parts of a lecture about '{title}':\n\n"
                        f"{all_parts}\n\n"
                        f"Write ONE unified 3-4 paragraph summary of the complete lecture. "
                        f"Flowing prose only — no headers or bullets. Output ONLY the summary:"
                    )
                    final_output = call_gemini("You are an expert academic summarizer.", merge_prompt, model, api_key)
                log_progress(f"Short-format merge complete — {word_count(final_output):,} words.")

            else:
                # Comprehensive / structured: full merge preserving everything
                merge_system = (
                    "You are an expert technical educator and document editor. "
                    "Your job is to merge multiple sets of lecture notes into ONE "
                    "perfect, comprehensive study document. "
                    "NEVER drop content. NEVER shorten anything. "
                    "The merged document must contain everything from every part."
                )

                total_notes_words = sum(word_count(o) for o in chunk_outputs)

                if total_notes_words <= 80_000:
                    final_output = call_gemini(
                        merge_system,
                        _build_merge_prompt(chunk_outputs, title, num_chunks),
                        model, api_key
                    )
                    log_progress(f"Direct merge complete — {word_count(final_output):,} words.")
                else:
                    log_progress(
                        f"Notes too large ({total_notes_words:,} words) for single merge. "
                        f"Using hierarchical merge…"
                    )
                    current_batch = chunk_outputs
                    pass_num      = 1

                    while len(current_batch) > 1:
                        next_batch = []
                        for i in range(0, len(current_batch), 2):
                            if i + 1 < len(current_batch):
                                pair       = [current_batch[i], current_batch[i+1]]
                                pair_label = f"Pass {pass_num}, chunks {i+1}-{i+2}"
                                log_progress(f"Merging {pair_label}…")
                                merged = call_gemini(
                                    merge_system,
                                    _build_merge_prompt(pair, title, 2),
                                    model, api_key
                                )
                                next_batch.append(merged)
                            else:
                                next_batch.append(current_batch[i])

                        current_batch = next_batch
                        pass_num     += 1
                        log_progress(
                            f"Merge pass {pass_num-1} done — {len(current_batch)} sections remain."
                        )

                    final_output = current_batch[0]
                    log_progress(f"Hierarchical merge complete — {word_count(final_output):,} words.")

        else:  # article mode
            fmt_config = FORMAT_MAP.get(tone_or_fmt, FORMAT_MAP["html"])
            combined   = "\n\n".join(chunk_outputs)
            unify_prompt = (
                f"You have {num_chunks} parts of an educational article about '{title}'.\n"
                f"Unify them into one seamless {fmt_config['desc']}:\n"
                f"- Write a proper introduction and conclusion\n"
                f"- Add a Table of Contents at the top\n"
                f"- Remove duplicated intro/outro sections from individual parts\n"
                f"- Keep ALL educational content, examples, and code\n"
                f"- Maintain consistent formatting throughout\n\n"
                f"INPUT PARTS:\n{combined[:100_000]}\n\n"
                f"Output the unified article now:"
            )
            final_output = strip_code_fences(call_gemini(
                "You are an expert technical content writer and editor.",
                unify_prompt,
                model, api_key
            ))

        log_progress("Merge complete.")

    output_wc = word_count(final_output)
    log_progress(f"Done! Generated {output_wc:,} words of output.")

    return {
        "content":    final_output,
        "chunks":     num_chunks,
        "words_in":   wc,
        "words_out":  output_wc,
        "model_used": model,
    }


# ═══════════════════════════════════════════════════════════════════════
# SECTION 5: BACKGROUND JOB SYSTEM
#
# For long transcripts (> ASYNC_THRESHOLD_WORDS), we run processing
# in a background daemon thread so the HTTP connection doesn't timeout.
#
# How it works:
#   1. POST /api/process/async → creates a job, starts a thread, returns job_id
#   2. Frontend polls GET /api/job/<job_id> every 3 seconds
#   3. Thread updates job state as each chunk completes
#   4. When done, job.status = "done" and job.result contains the output
#
# Job storage: in-memory dict (fine for personal/single-user use)
# For multi-user production: replace with Redis or a database
# ═══════════════════════════════════════════════════════════════════════

# Global job store — dict keyed by job_id UUID strings
# Each value is a dict with: status, progress, result, error, etc.
_jobs: dict      = {}
_jobs_lock       = threading.Lock()  # thread-safe access to _jobs


def _create_job() -> str:
    """
    Create a new job entry and return its unique ID.
    The job starts in 'queued' status.
    Also cleans up stale jobs older than 4 hours to prevent memory leaks.
    """
    job_id = str(uuid.uuid4())

    with _jobs_lock:
        # ── Purge stale jobs (older than 4 hours) ──────────────
        now   = time.time()
        stale = [
            jid for jid, job in _jobs.items()
            if now - job.get("created_at", 0) > 14400  # 4 hours
        ]
        for jid in stale:
            del _jobs[jid]
        if stale:
            logger.info("Purged %d stale jobs from memory", len(stale))

        _jobs[job_id] = {
            "status":       "queued",
            "progress":     [],
            "result":       None,
            "error":        None,
            "created_at":   now,
            "model_used":   "",
            "chunks_total": 0,
            "chunks_done":  0,
            "percent":      0,
        }
    return job_id


def _update_job(job_id: str, **fields):
    """Thread-safe update of job fields."""
    with _jobs_lock:
        if job_id in _jobs:
            _jobs[job_id].update(fields)


def _add_progress(job_id: str, message: str):
    """
    Append a timestamped progress message to the job log.
    These messages are shown in the frontend's live progress terminal.
    """
    timestamped = f"[{time.strftime('%H:%M:%S')}] {message}"
    with _jobs_lock:
        if job_id in _jobs:
            _jobs[job_id]["progress"].append(timestamped)
    logger.info("[job %s] %s", job_id[:8], message)


def _run_job_in_background(
    job_id:     str,
    transcript: str,
    mode:       str,
    fmt:        str,
    title:      str,
    model:      str,
    api_key:    str = "",
):
    """
    Background thread function — runs the full processing pipeline
    and updates job state at each step.
    """
    try:
        _update_job(job_id, status="running", model_used=model)
        _add_progress(job_id, f"Starting {mode} generation with {model}")

        wc         = word_count(transcript)
        num_chunks = math.ceil(wc / CHUNK_WORDS)
        _update_job(job_id, chunks_total=num_chunks)
        _add_progress(job_id, f"Transcript: {wc:,} words → {num_chunks} chunk(s) to process")

        chunks_completed = [0]

        def on_progress(msg: str):
            _add_progress(job_id, msg)
            if "chunk" in msg.lower() and "complete" in msg.lower():
                chunks_completed[0] += 1
                pct = min(int((chunks_completed[0] / num_chunks) * 85), 85)
                _update_job(job_id, chunks_done=chunks_completed[0], percent=pct)
            elif "merge" in msg.lower() and "complete" in msg.lower():
                _update_job(job_id, percent=95)
            elif msg.startswith("Done!"):
                _update_job(job_id, percent=100)

        result = process_transcript(
            transcript  = transcript,
            mode        = mode,
            tone_or_fmt = fmt,
            title       = title,
            model       = model,
            api_key     = api_key,
            progress_cb = on_progress,
        )

        result["model_used"] = model
        _update_job(job_id, status="done", result=result, percent=100)
        _add_progress(job_id, f"✓ Job complete — {result['words_out']:,} words output")

    except Exception as exc:
        logger.exception("[job %s] Failed with exception", job_id[:8])
        _add_progress(job_id, f"ERROR: {exc}")
        _update_job(job_id, status="error", error=str(exc))


# ═══════════════════════════════════════════════════════════════════════
# SECTION 6: DOCUMENT EXPORT (PDF and DOCX)
# ═══════════════════════════════════════════════════════════════════════

# Base HTML template used when exporting to HTML or PDF
HTML_EXPORT_TEMPLATE = """<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<meta name="viewport" content="width=device-width, initial-scale=1.0">
<title>{title}</title>
<style>
  /* Clean, readable typography for exported documents */
  body {{
    font-family: Georgia, 'Times New Roman', serif;
    max-width: 800px;
    margin: 3rem auto;
    padding: 0 2rem;
    line-height: 1.8;
    color: #1a1a1a;
    font-size: 16px;
  }}
  h1 {{ font-size: 2rem; margin-bottom: 0.4rem; line-height: 1.2; border-bottom: 2px solid #333; padding-bottom: 0.5rem; }}
  h2 {{ font-size: 1.4rem; margin: 2rem 0 0.6rem; color: #222; border-bottom: 1px solid #ddd; padding-bottom: 0.3rem; }}
  h3 {{ font-size: 1.1rem; margin: 1.5rem 0 0.4rem; color: #333; }}
  h4 {{ font-size: 0.95rem; margin: 1rem 0 0.3rem; color: #444; font-weight: 600; }}
  .meta {{ color: #888; font-size: 0.82rem; font-family: monospace; margin-bottom: 2rem;
           padding-bottom: 1rem; border-bottom: 1px solid #e5e5e5; }}
  p  {{ margin-bottom: 0.9rem; }}
  ul, ol {{ padding-left: 1.6rem; margin: 0.5rem 0 1rem; }}
  li {{ margin-bottom: 0.35rem; }}
  blockquote {{ border-left: 4px solid #6c63ff; padding: 0.5rem 1.25rem; margin: 1rem 0;
               color: #555; font-style: italic; background: #f4f3ff; border-radius: 0 4px 4px 0; }}
  code {{ background: #f3f4f6; font-family: 'Courier New', monospace; font-size: 0.85em;
          padding: 0.15em 0.4em; border-radius: 3px; }}
  pre  {{ background: #1e1e2e; color: #cdd6f4; padding: 1.25rem; border-radius: 6px;
          overflow-x: auto; margin: 1rem 0; font-size: 0.85rem; line-height: 1.5; }}
  pre code {{ background: none; color: inherit; padding: 0; font-size: inherit; }}
  dt {{ font-weight: 700; margin-top: 0.6rem; }}
  dd {{ margin-left: 1.5rem; color: #444; margin-bottom: 0.3rem; }}
  table {{ border-collapse: collapse; width: 100%; margin: 1rem 0; font-size: 0.9rem; }}
  th, td {{ border: 1px solid #ddd; padding: 0.5rem 0.75rem; text-align: left; }}
  th {{ background: #f0f0f0; font-weight: 700; }}
  strong {{ font-weight: 700; color: #111; }}
  @media print {{
    body {{ max-width: 100%; margin: 0; padding: 1rem; }}
    pre {{ white-space: pre-wrap; }}
  }}
</style>
</head>
<body>
{body}
</body>
</html>"""


def generate_pdf_bytes(html_content: str, title: str = "Document") -> bytes:
    """
    Convert HTML or Markdown content to a PDF using ReportLab.

    ReportLab is pure Python with no system dependencies (no cairo, no pango).
    It is already installed as a dependency of xhtml2pdf, or install directly:
        pip install reportlab

    Parses the HTML/Markdown into styled paragraphs and renders to PDF.
    Supports: headings (h1-h4), paragraphs, bullet lists, numbered lists,
              code blocks, blockquotes, bold/italic inline text, hr.

    Parameters:
        html_content — HTML body content or Markdown (auto-detected)
        title        — document title shown at the top

    Returns: PDF file as bytes
    Raises: RuntimeError if reportlab is not installed
    """
    if not PDF_AVAILABLE:
        raise RuntimeError(
            "reportlab not installed. Run: pip install reportlab\n"
            "This is a pure Python library — no system dependencies needed.\n"
            "It works on Mac, Windows, and Linux without any C libraries."
        )

    # ── Detect format and normalise to plain Markdown ─────────────────
    # If it looks like HTML, strip tags first to get clean text for parsing
    content = html_content.strip()
    if content.startswith("<!DOCTYPE") or content.startswith("<html"):
        # Full HTML document — extract body
        body_match = re.search(r"<body[^>]*>(.*?)</body>", content, re.DOTALL | re.IGNORECASE)
        content = body_match.group(1) if body_match else content

    # Convert HTML tags to Markdown-ish lines for uniform parsing
    if re.search(r"<h[1-6]|<p |<ul|<ol|<li|<pre|<blockquote", content, re.IGNORECASE):
        content = html_to_markdown(content)

    # ── Build ReportLab styles ─────────────────────────────────────────
    base = getSampleStyleSheet()

    PAGE_W, PAGE_H = A4
    MARGIN = 2.2 * cm

    def _style(name, parent="Normal", **kwargs):
        s = ParagraphStyle(name, parent=base[parent], **kwargs)
        return s

    DARK   = rl_colors.HexColor("#1a1a2e")
    ACCENT = rl_colors.HexColor("#6c63ff")
    GREY   = rl_colors.HexColor("#555555")
    CODE_BG = rl_colors.HexColor("#f3f4f6")

    styles = {
        "h1": _style("H1", "Heading1",
            fontSize=22, textColor=DARK, spaceAfter=10, spaceBefore=18,
            fontName="Helvetica-Bold", leading=28),
        "h2": _style("H2", "Heading2",
            fontSize=16, textColor=DARK, spaceAfter=6, spaceBefore=14,
            fontName="Helvetica-Bold", leading=22,
            borderPadding=(0, 0, 4, 0)),
        "h3": _style("H3", "Heading3",
            fontSize=13, textColor=ACCENT, spaceAfter=4, spaceBefore=10,
            fontName="Helvetica-Bold", leading=18),
        "h4": _style("H4", "Heading4",
            fontSize=11, textColor=GREY, spaceAfter=3, spaceBefore=8,
            fontName="Helvetica-BoldOblique", leading=15),
        "body": _style("Body", "Normal",
            fontSize=10.5, textColor=DARK, spaceAfter=5, spaceBefore=2,
            leading=16, fontName="Helvetica"),
        "bullet": _style("Bullet", "Normal",
            fontSize=10.5, textColor=DARK, spaceAfter=3, spaceBefore=1,
            leading=15, fontName="Helvetica",
            leftIndent=18, bulletIndent=6),
        "blockquote": _style("BQ", "Normal",
            fontSize=10.5, textColor=GREY, spaceAfter=6, spaceBefore=6,
            leading=16, fontName="Helvetica-Oblique",
            leftIndent=20, rightIndent=10,
            borderColor=ACCENT, borderPadding=(4, 4, 4, 10), borderWidth=0),
        "code": _style("Code", "Code",
            fontSize=8.5, fontName="Courier",
            textColor=rl_colors.HexColor("#2d2d2d"),
            backColor=CODE_BG,
            spaceAfter=8, spaceBefore=6,
            leading=13, leftIndent=10, rightIndent=10),
        "meta": _style("Meta", "Normal",
            fontSize=9, textColor=GREY, spaceAfter=14, spaceBefore=0,
            fontName="Helvetica-Oblique"),
        "hr": _style("HR", "Normal", spaceAfter=8, spaceBefore=8),
    }

    # ── Sanitise text for ReportLab XML parser ─────────────────────────
    def _rl_escape(text: str) -> str:
        """Escape XML special chars and convert simple Markdown bold/italic/code."""
        # XML-safe first
        text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        # Bold **text**
        text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
        # Italic *text*
        text = re.sub(r"\*(.+?)\*", r"<i>\1</i>", text)
        # Inline code `text`
        text = re.sub(r"`([^`]+)`", r'<font name="Courier" size="9">\1</font>', text)
        return text

    # ── Parse Markdown lines into ReportLab flowables ──────────────────
    story = []
    lines = content.splitlines()
    i = 0

    # Document title at top
    story.append(Paragraph(_rl_escape(title), styles["h1"]))
    story.append(HRFlowable(width="100%", thickness=1.5, color=ACCENT, spaceAfter=6))
    story.append(Paragraph(
        f"Generated by LectureAI · {time.strftime('%B %d, %Y')}",
        styles["meta"]
    ))

    while i < len(lines):
        line = lines[i]
        raw  = line.strip()

        # ── Headings ──────────────────────────────────────────────────
        if raw.startswith("#### "):
            story.append(Paragraph(_rl_escape(raw[5:]), styles["h4"]))

        elif raw.startswith("### "):
            story.append(Paragraph(_rl_escape(raw[4:]), styles["h3"]))

        elif raw.startswith("## "):
            story.append(Spacer(1, 4))
            story.append(Paragraph(_rl_escape(raw[3:]), styles["h2"]))
            story.append(HRFlowable(width="100%", thickness=0.5,
                                    color=rl_colors.HexColor("#dddddd"), spaceAfter=2))

        elif raw.startswith("# "):
            story.append(Paragraph(_rl_escape(raw[2:]), styles["h1"]))

        # ── Horizontal rule ───────────────────────────────────────────
        elif raw in ("---", "***", "___"):
            story.append(HRFlowable(width="100%", thickness=0.7,
                                    color=rl_colors.HexColor("#cccccc"),
                                    spaceBefore=6, spaceAfter=6))

        # ── Code blocks ───────────────────────────────────────────────
        elif raw.startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            code_text = "\n".join(code_lines)
            # Preformatted respects whitespace exactly
            story.append(Preformatted(code_text, styles["code"]))

        # ── Blockquotes ───────────────────────────────────────────────
        elif raw.startswith("> "):
            story.append(Paragraph(
                _rl_escape(raw[2:]),
                styles["blockquote"]
            ))

        # ── Bullet list items ─────────────────────────────────────────
        elif re.match(r"^[\*\-] ", raw):
            story.append(Paragraph(
                "• " + _rl_escape(raw[2:]),
                styles["bullet"]
            ))

        # ── Numbered list items ───────────────────────────────────────
        elif re.match(r"^\d+\. ", raw):
            num   = re.match(r"^(\d+)\. ", raw).group(1)
            text  = re.sub(r"^\d+\. ", "", raw)
            story.append(Paragraph(
                f"<b>{num}.</b> " + _rl_escape(text),
                styles["bullet"]
            ))

        # ── Regular paragraph text ────────────────────────────────────
        elif raw:
            story.append(Paragraph(_rl_escape(raw), styles["body"]))

        # ── Blank line → small spacer ────────────────────────────────
        else:
            story.append(Spacer(1, 3))

        i += 1

    # ── Render to bytes ────────────────────────────────────────────────
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize    = A4,
        leftMargin  = MARGIN,
        rightMargin = MARGIN,
        topMargin   = MARGIN,
        bottomMargin= MARGIN,
        title       = title,
        author      = "LectureAI",
    )

    def _add_page_number(canvas_obj, doc_obj):
        """Draw page number at bottom centre of each page."""
        canvas_obj.saveState()
        canvas_obj.setFont("Helvetica", 8)
        canvas_obj.setFillColor(GREY)
        canvas_obj.drawCentredString(
            PAGE_W / 2, 1.2 * cm,
            f"Page {doc_obj.page}  ·  {title[:60]}"
        )
        canvas_obj.restoreState()

    doc.build(story, onFirstPage=_add_page_number, onLaterPages=_add_page_number)
    buffer.seek(0)
    return buffer.read()


def generate_docx_bytes(content: str, title: str = "Document", is_markdown: bool = True) -> bytes:
    """
    Convert Markdown or HTML content to a Word DOCX file.

    Uses python-docx. Pt, RGBColor, Inches are imported at module level
    so _add_inline_formatting can access them without NameError.

    Parameters:
        content     — the text content (Markdown or HTML)
        title       — document title (used as the first heading)
        is_markdown — True if content is Markdown, False if HTML

    Returns: DOCX file as bytes
    Raises: RuntimeError if python-docx is not installed
    """
    if not DOCX_AVAILABLE:
        raise RuntimeError(
            "python-docx not installed. Run: pip install python-docx"
        )

    # If content is HTML, convert to Markdown first (easier to parse into docx)
    if not is_markdown:
        content = html_to_markdown(content)

    doc = DocxDocument()

    # ── Document-level styles ──────────────────────────────────────
    # Set the default body font
    normal_style           = doc.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)

    # Set reasonable page margins
    for section in doc.sections:
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin   = Inches(1.25)
        section.right_margin  = Inches(1.25)

    # ── Parse Markdown line by line and add to document ────────────
    lines   = content.splitlines()
    i       = 0
    in_list = False   # track whether we're inside a list block

    while i < len(lines):
        line = lines[i]

        # ── Headings ──────────────────────────────────────────────
        if line.startswith("#### "):
            in_list = False
            doc.add_heading(line[5:].strip(), level=4)

        elif line.startswith("### "):
            in_list = False
            doc.add_heading(line[4:].strip(), level=3)

        elif line.startswith("## "):
            in_list = False
            doc.add_heading(line[3:].strip(), level=2)

        elif line.startswith("# "):
            in_list = False
            doc.add_heading(line[2:].strip(), level=1)

        # ── Fenced code blocks ────────────────────────────────────
        elif line.startswith("```"):
            in_list    = False
            code_lines = []
            i         += 1
            # Collect lines until closing ```
            while i < len(lines) and not lines[i].startswith("```"):
                code_lines.append(lines[i])
                i += 1
            # Add as monospace paragraph
            if code_lines:
                para = doc.add_paragraph("\n".join(code_lines))
                # "No Spacing" removes extra paragraph spacing
                try:
                    para.style = doc.styles["No Spacing"]
                except KeyError:
                    pass  # style may not exist in all Word templates
                for run in para.runs:
                    run.font.name  = "Courier New"
                    run.font.size  = Pt(9)
                    run.font.color.rgb = RGBColor(0x1E, 0x1E, 0x2E)

        # ── Blockquotes ───────────────────────────────────────────
        elif line.startswith("> "):
            in_list = False
            para    = doc.add_paragraph(line[2:].strip())
            para.paragraph_format.left_indent = Inches(0.4)
            for run in para.runs:
                run.italic = True

        # ── Bullet lists ──────────────────────────────────────────
        elif re.match(r"^[\*\-] ", line):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")
            in_list = True

        # ── Numbered lists ────────────────────────────────────────
        elif re.match(r"^\d+\. ", line):
            text = re.sub(r"^\d+\. ", "", line).strip()
            doc.add_paragraph(text, style="List Number")
            in_list = True

        # ── Horizontal rules ──────────────────────────────────────
        elif line.strip() in ("---", "***", "___"):
            doc.add_paragraph("─" * 60)

        # ── Regular text (with inline bold/italic/code support) ───
        elif line.strip():
            in_list = False
            para    = doc.add_paragraph()
            _add_inline_formatting(para, line.strip())

        # ── Blank lines — just reset list state ───────────────────
        else:
            in_list = False

        i += 1

    # Save to an in-memory buffer (no temp files needed)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


def _add_inline_formatting(paragraph, text: str):
    """
    Add text to a docx paragraph with inline Markdown formatting applied.

    Handles: **bold**, *italic*, `inline code`

    Pt, RGBColor are imported at module level (not inside this function)
    so there is no NameError even when called from deeply nested code.
    """
    # Split text on bold (**text**), italic (*text*), and code (`text`) markers
    pattern = r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)"
    parts   = re.split(pattern, text)

    for part in parts:
        if not part:
            continue

        if part.startswith("**") and part.endswith("**") and len(part) > 4:
            run      = paragraph.add_run(part[2:-2])
            run.bold = True

        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            run        = paragraph.add_run(part[1:-1])
            run.italic = True

        elif part.startswith("`") and part.endswith("`") and len(part) > 2:
            run           = paragraph.add_run(part[1:-1])
            run.font.name = "Courier New"
            # Pt is imported at module level — safe to use here
            if Pt is not None:
                run.font.size = Pt(9)

        else:
            paragraph.add_run(part)


def html_to_markdown(html_str: str) -> str:
    """
    Convert HTML to Markdown for DOCX export.

    This is a simple regex-based conversion that handles the most common
    HTML elements produced by Gemini. It's not a full HTML parser —
    for complex HTML you would use a library like html2text.
    """
    md = html_str

    # Convert headings
    md = re.sub(r"<h1[^>]*>(.*?)</h1>", r"# \1\n",        md, flags=re.DOTALL | re.IGNORECASE)
    md = re.sub(r"<h2[^>]*>(.*?)</h2>", r"## \1\n",       md, flags=re.DOTALL | re.IGNORECASE)
    md = re.sub(r"<h3[^>]*>(.*?)</h3>", r"### \1\n",      md, flags=re.DOTALL | re.IGNORECASE)
    md = re.sub(r"<h4[^>]*>(.*?)</h4>", r"#### \1\n",     md, flags=re.DOTALL | re.IGNORECASE)

    # Convert blockquotes
    md = re.sub(r"<blockquote[^>]*>(.*?)</blockquote>",
                lambda m: "> " + m.group(1).strip() + "\n",
                md, flags=re.DOTALL | re.IGNORECASE)

    # Convert lists
    md = re.sub(r"<li[^>]*>(.*?)</li>", r"- \1\n",        md, flags=re.DOTALL | re.IGNORECASE)

    # Convert code blocks
    md = re.sub(r"<pre[^>]*><code[^>]*>(.*?)</code></pre>",
                lambda m: "```\n" + html.unescape(m.group(1).strip()) + "\n```\n",
                md, flags=re.DOTALL | re.IGNORECASE)

    # Convert inline code
    md = re.sub(r"<code[^>]*>(.*?)</code>", r"`\1`",       md, flags=re.DOTALL | re.IGNORECASE)

    # Convert inline formatting
    md = re.sub(r"<strong[^>]*>(.*?)</strong>", r"**\1**", md, flags=re.DOTALL | re.IGNORECASE)
    md = re.sub(r"<em[^>]*>(.*?)</em>",         r"*\1*",   md, flags=re.DOTALL | re.IGNORECASE)
    md = re.sub(r"<b[^>]*>(.*?)</b>",           r"**\1**", md, flags=re.DOTALL | re.IGNORECASE)

    # Convert definition lists
    md = re.sub(r"<dt[^>]*>(.*?)</dt>", r"**\1**\n",       md, flags=re.DOTALL | re.IGNORECASE)
    md = re.sub(r"<dd[^>]*>(.*?)</dd>", r"  \1\n",         md, flags=re.DOTALL | re.IGNORECASE)

    # Convert paragraphs to double newlines
    md = re.sub(r"<p[^>]*>(.*?)</p>", r"\1\n\n",           md, flags=re.DOTALL | re.IGNORECASE)

    # Strip all remaining HTML tags
    md = re.sub(r"<[^>]+>", "", md)

    # Decode HTML entities
    md = html.unescape(md)

    # Clean up excessive blank lines
    md = re.sub(r"\n{3,}", "\n\n", md)

    return md.strip()


# ═══════════════════════════════════════════════════════════════════════
# SECTION 7: PROMPT TEMPLATES
# Maps for tone (note styles) and format (article styles)
# ═══════════════════════════════════════════════════════════════════════

TONE_MAP = {
    "comprehensive": (
        "Write COMPREHENSIVE lecture notes that capture every concept, algorithm, formula, "
        "code snippet, example, and insight. Use clear headings and subheadings. "
        "A student should be able to learn everything from these notes alone — miss nothing."
    ),
    "concise": (
        "Write a concise 3-5 sentence summary capturing the core message and key points only."
    ),
    "detailed": (
        "Write a detailed 3-4 paragraph summary covering all major themes, arguments, "
        "techniques, and key takeaways."
    ),
    "notes": (
        "Write structured lecture notes with clear headings, sub-points, formulas, "
        "code examples, and key definitions. Capture everything the instructor covers."
    ),
    "bullet": (
        "Write a bullet-point summary with 8-15 key takeaways covering the most important "
        "facts, concepts, algorithms, formulas, and practical insights."
    ),
    "executive": (
        "Write a 3-sentence executive summary: one sentence for the core topic, "
        "one for the key insight or technique, one for the main actionable conclusion."
    ),
}

FORMAT_MAP = {
    "html": {
        "desc": "a comprehensive HTML educational article",
        "instructions": (
            "- <h1> for the lecture/course title\n"
            '- <p class="meta"> with byline and estimated read time\n'
            "- <h2> for each major topic\n"
            "- <h3> for subtopics with full <p> explanations\n"
            "- <ul>/<ol> for step-by-step processes and key points\n"
            "- <blockquote> for key definitions and insights\n"
            "- <pre><code> for all code examples\n"
            "- Final <h2>Key Takeaways</h2> with comprehensive <ul>\n"
            "Return ONLY inner HTML starting with <h1>. No <html>/<body>/code fences."
        ),
    },
    "markdown": {
        "desc": "comprehensive Markdown study notes",
        "instructions": (
            "- # Course/lecture title\n"
            "- *Metadata line*\n"
            "- ## Major topics with full explanations\n"
            "- ### Subtopics\n"
            "- ``` code blocks for all code\n"
            "- > blockquotes for key definitions\n"
            "- ## Key Takeaways at the end\n"
            "Return ONLY Markdown. No code fences wrapping the whole output."
        ),
    },
    "blog": {
        "desc": "an in-depth educational blog post (HTML)",
        "instructions": (
            "- Engaging <h1> headline\n"
            "- Compelling hook introduction paragraph\n"
            "- <h2> sections for each concept with full technical explanations\n"
            "- Code in <pre><code> blocks\n"
            "- <h2>What You've Learned</h2> conclusion\n"
            "Return ONLY inner HTML starting with <h1>."
        ),
    },
    "studynotes": {
        "desc": "complete HTML study guide",
        "instructions": (
            "- <h1> course/topic title\n"
            '- <p class="meta"> subject and estimated study time\n'
            "- <h2>Prerequisites</h2>\n"
            "- <h2>Overview</h2> — 2-3 sentence synopsis\n"
            "- Multiple <h2> topic sections, each with:\n"
            "    <h3> subtopics, <p> explanations, <ul> key points, <pre><code> for code\n"
            "- <h2>Key Definitions</h2> as <dl><dt>/<dd>\n"
            "- <h2>Review Questions</h2> as <ol> with 8-10 questions\n"
            "- <h2>Further Reading / Resources</h2>\n"
            "Return ONLY inner HTML starting with <h1>."
        ),
    },
    "newsletter": {
        "desc": "a technical newsletter edition (HTML)",
        "instructions": (
            "- Bold <h1> subject-line headline\n"
            '- <p class="meta"> with edition label\n'
            "- Brief 2-sentence intro\n"
            "- <h2>What Was Covered</h2>\n"
            "- <h2>Key Concepts Explained</h2> — detailed explanations\n"
            "- <h2>Practical Takeaways</h2> — <ul> with actionable items\n"
            "Return ONLY inner HTML starting with <h1>."
        ),
    },
}


# ═══════════════════════════════════════════════════════════════════════
# SECTION 8: FLASK ROUTES
# ═══════════════════════════════════════════════════════════════════════

@app.route("/")
def index():
    """Serve the main application page."""
    return render_template("index.html", model=GEMINI_MODEL)


@app.route("/health")
def health():
    """
    Health check endpoint for load balancers and uptime monitors.
    Returns 200 OK with basic server info.
    """
    return jsonify({
        "status":       "ok",
        "model":        GEMINI_MODEL,
        "async_threshold_words": ASYNC_THRESHOLD_WORDS,
    })


@app.route("/api/config")
def api_config():
    """
    Return server configuration to the frontend.
    Called on page load so the frontend knows which model is active
    and what the chunking thresholds are.
    """
    return jsonify({
        "model":                  GEMINI_MODEL,
        "max_input_k":            round(SAFE_INPUT_WORDS / 1000),
        "chunk_words":            CHUNK_WORDS,
        "async_threshold_words":  ASYNC_THRESHOLD_WORDS,
        "has_server_key":         bool(GEMINI_API_KEY),
        "has_groq_key":           bool(GROQ_API_KEY),
        "max_retries":            GEMINI_MAX_RETRIES,
        "base_wait_secs":         GEMINI_BASE_WAIT,
    })


@app.route("/api/transcript", methods=["POST"])
def api_transcript():
    """
    Fetch the transcript for a YouTube video.

    Request body:
        { "url": "https://youtube.com/watch?v=..." }

    Response:
        {
            "transcript":       str,  — full transcript text
            "title":            str,  — video title
            "channel":          str,  — channel name
            "word_count":       int,  — number of words
            "estimated_hours":  float,— estimated video duration
            "chunks_needed":    int,  — how many processing chunks needed
            "method":           str   — which fetch method succeeded
        }

    Error responses:
        400 — missing or invalid URL
        422 — transcript fetch failed (with error message)
        500 — unexpected server error
    """
    data     = request.get_json(silent=True) or {}
    url      = (data.get("url") or "").strip()

    if not url:
        return jsonify({"error": "Missing 'url' field in request body."}), 400

    video_id = extract_video_id(url)
    if not video_id:
        return jsonify({"error": "Invalid YouTube URL — could not extract video ID."}), 400

    try:
        result = fetch_transcript(video_id)
        return jsonify(result)
    except ValueError as e:
        # Expected error (e.g., no captions) — return 422 with user-friendly message
        logger.warning("Transcript fetch failed for %s: %s", video_id, e)
        return jsonify({"error": str(e)}), 422
    except Exception as e:
        # Unexpected error — log full traceback and return 500
        logger.exception("Unexpected error fetching transcript for %s", video_id)
        return jsonify({"error": f"Server error: {e}"}), 500


@app.route("/api/process", methods=["POST"])
def api_process():
    """
    Process a transcript synchronously (waits for completion).
    Use for short transcripts only (< ASYNC_THRESHOLD_WORDS words).
    For long transcripts, use /api/process/async instead.

    Request body:
        {
            "transcript": str,    — the transcript text
            "mode":       str,    — "notes" or "article"
            "format":     str,    — tone/format key
            "title":      str,    — video title (optional)
            "model":      str     — model name (optional, overrides env default)
        }

    Response:
        {
            "content":    str,  — generated notes or article
            "chunks":     int,  — number of chunks processed
            "words_in":   int,  — input word count
            "words_out":  int,  — output word count
            "model_used": str   — which model was used
        }
    """
    data       = request.get_json(silent=True) or {}
    transcript = (data.get("transcript") or "").strip()
    mode       = (data.get("mode") or "notes").strip().lower()
    fmt        = (data.get("format") or "comprehensive").strip().lower()
    title      = (data.get("title") or "").strip()
    model_req  = (data.get("model") or "").strip() or GEMINI_MODEL
    api_key    = (data.get("api_key") or "").strip()

    # Validate inputs
    if not transcript:
        return jsonify({"error": "Missing 'transcript' field."}), 400
    if mode not in ("notes", "article"):
        return jsonify({"error": "Invalid 'mode' — must be 'notes' or 'article'."}), 400

    try:
        result = process_transcript(
            transcript   = transcript,
            mode         = mode,
            tone_or_fmt  = fmt,
            title        = title,
            model        = model_req,
            api_key      = api_key,
        )
        result["model_used"] = model_req
        return jsonify(result)

    except (ValueError, RuntimeError, TimeoutError, ConnectionError) as e:
        logger.error("Process error: %s", e)
        return jsonify({"error": str(e)}), 502
    except Exception as e:
        logger.exception("Unexpected process error")
        return jsonify({"error": f"Server error: {e}"}), 500


@app.route("/api/process/async", methods=["POST"])
def api_process_async():
    """
    Start a background processing job and return immediately.
    The frontend should poll GET /api/job/<job_id> for progress.

    Use this for transcripts > ASYNC_THRESHOLD_WORDS words to avoid
    browser HTTP timeouts on very long processing jobs.

    Request body: same as /api/process

    Response:
        {
            "job_id":        str,  — UUID to poll for status
            "words":         int,  — transcript word count
            "chunks_total":  int,  — total chunks to process
            "estimated_min": int,  — estimated minimum minutes
            "estimated_max": int   — estimated maximum minutes
        }
    """
    data       = request.get_json(silent=True) or {}
    transcript = (data.get("transcript") or "").strip()
    mode       = (data.get("mode") or "notes").strip().lower()
    fmt        = (data.get("format") or "comprehensive").strip().lower()
    title      = (data.get("title") or "").strip()
    model_req  = (data.get("model") or "").strip() or GEMINI_MODEL
    api_key    = (data.get("api_key") or "").strip()

    if not transcript:
        return jsonify({"error": "Missing 'transcript'."}), 400
    if mode not in ("notes", "article"):
        return jsonify({"error": "mode must be 'notes' or 'article'."}), 400

    # Create the job record and start the background thread
    job_id     = _create_job()
    wc         = word_count(transcript)
    num_chunks = math.ceil(wc / CHUNK_WORDS)

    # Estimate seconds per chunk based on model (flash is ~2x faster than pro)
    is_flash   = "flash" in model_req.lower()
    secs_chunk = 45 if is_flash else 90
    est_min    = max(1, round(num_chunks * secs_chunk / 60))

    thread = threading.Thread(
        target   = _run_job_in_background,
        args     = (job_id, transcript, mode, fmt, title, model_req, api_key),
        daemon   = True,
        name     = f"job-{job_id[:8]}"
    )
    thread.start()

    logger.info(
        "Async job started: id=%s words=%d chunks=%d model=%s",
        job_id[:8], wc, num_chunks, model_req
    )

    return jsonify({
        "job_id":        job_id,
        "words":         wc,
        "chunks_total":  num_chunks,
        "estimated_min": est_min,
        "estimated_max": est_min * 2,
    })


@app.route("/api/job/<job_id>", methods=["GET"])
def api_job_status(job_id: str):
    """
    Get the current status and progress of a background job.

    Called by the frontend every 3 seconds to update the progress modal.

    Response:
        {
            "job_id":       str,
            "status":       "queued" | "running" | "done" | "error",
            "percent":      int,    — 0-100 progress percentage
            "progress":     list,   — timestamped log messages
            "chunks_total": int,
            "chunks_done":  int,
            "model_used":   str,
            "result":       dict,   — only present when status = "done"
            "error":        str     — only present when status = "error"
        }
    """
    with _jobs_lock:
        job = _jobs.get(job_id)

    if not job:
        return jsonify({"error": f"Job '{job_id}' not found."}), 404

    # Build response — always include status/progress fields
    response = {
        "job_id":       job_id,
        "status":       job["status"],
        "percent":      job["percent"],
        "progress":     job["progress"],
        "chunks_total": job["chunks_total"],
        "chunks_done":  job["chunks_done"],
        "model_used":   job["model_used"],
    }

    # Add result only when done (avoids sending large data on every poll)
    if job["status"] == "done":
        response["result"] = job["result"]
    if job["status"] == "error":
        response["error"] = job["error"]

    return jsonify(response)


@app.route("/api/job/<job_id>", methods=["DELETE"])
def api_job_delete(job_id: str):
    """
    Remove a completed job from memory to free up space.
    Called by the frontend after it has received the final result.
    """
    with _jobs_lock:
        removed = _jobs.pop(job_id, None)
    return jsonify({"deleted": removed is not None})


@app.route("/api/summarize", methods=["POST"])
def api_summarize():
    """
    Summarize a transcript using the configured tone style.
    Legacy endpoint kept for backward compatibility.

    Request body:
        {
            "transcript": str,
            "tone":       str,  — key from TONE_MAP
            "model":      str   — optional model override
        }
    """
    data       = request.get_json(silent=True) or {}
    transcript = (data.get("transcript") or "").strip()
    tone       = (data.get("tone") or "concise").strip().lower()
    model_req  = (data.get("model") or "").strip() or GEMINI_MODEL
    api_key    = (data.get("api_key") or "").strip()

    if not transcript:
        return jsonify({"error": "Missing 'transcript'."}), 400
    if tone not in TONE_MAP:
        return jsonify({"error": f"Invalid tone. Options: {', '.join(TONE_MAP)}."}), 400

    system = (
        "You are an expert academic summarizer. Convert lecture transcripts into clear, "
        "accurate written content. Never use phrases like 'In this video' or 'The speaker says'."
    )
    prompt = (
        f"STYLE: {TONE_MAP[tone]}\n\n"
        f"TRANSCRIPT:\n{transcript}\n\n"
        f"Output ONLY the summary — no preamble, labels, or meta-commentary."
    )

    try:
        summary = call_gemini(system, prompt, model_req, api_key)
    except Exception as e:
        logger.error("Summarize error: %s", e)
        return jsonify({"error": str(e)}), 502

    wc_in  = word_count(transcript)
    wc_out = word_count(summary)

    return jsonify({
        "summary":        summary,
        "word_count_in":  wc_in,
        "word_count_out": wc_out,
        "reduction_pct":  round((1 - wc_out / max(wc_in, 1)) * 100),
        "tokens_est":     round(wc_in * 1.33),
        "model_used":     model_req,
    })


@app.route("/api/article", methods=["POST"])
def api_article():
    """
    Generate a formatted article from a transcript and summary.

    Request body:
        {
            "summary":    str,  — previously generated summary
            "transcript": str,  — original transcript for context
            "format":     str,  — key from FORMAT_MAP
            "model":      str   — optional model override
        }
    """
    data       = request.get_json(silent=True) or {}
    summary    = (data.get("summary") or "").strip()
    transcript = (data.get("transcript") or "").strip()
    fmt        = (data.get("format") or "html").strip().lower()
    model_req  = (data.get("model") or "").strip() or GEMINI_MODEL
    api_key    = (data.get("api_key") or "").strip()

    if not summary:
        return jsonify({"error": "Missing 'summary'."}), 400
    if not transcript:
        return jsonify({"error": "Missing 'transcript'."}), 400
    if fmt not in FORMAT_MAP:
        return jsonify({"error": f"Invalid format. Options: {', '.join(FORMAT_MAP)}."}), 400

    fc = FORMAT_MAP[fmt]
    system = (
        "You are an expert technical content writer. Transform lecture transcripts into "
        "polished, publication-ready educational content. Capture everything — miss nothing. "
        "Expand meaningfully on ideas — never just rephrase."
    )
    prompt = (
        f"Create {fc['desc']} from this lecture.\n\n"
        f"SUMMARY:\n{summary}\n\n"
        f"TRANSCRIPT (extra context):\n{transcript[:60_000]}\n\n"
        f"FORMAT INSTRUCTIONS:\n{fc['instructions']}"
    )

    try:
        article = strip_code_fences(call_gemini(system, prompt, model_req, api_key))
    except Exception as e:
        logger.error("Article error: %s", e)
        return jsonify({"error": str(e)}), 502

    return jsonify({"article": article, "format": fmt, "model_used": model_req})


@app.route("/api/verify", methods=["POST"])
def api_verify():
    """
    Coverage audit — verifies notes fully cover the transcript.

    Bug fixed: previously only checked first 30k words of notes and 40k of transcript.
    Now checks the full transcript (sampled at even intervals) and full notes.
    """
    data       = request.get_json(silent=True) or {}
    transcript = (data.get("transcript") or "").strip()
    notes      = (data.get("notes") or "").strip()
    model_req  = (data.get("model") or "").strip() or GEMINI_MODEL
    api_key    = (data.get("api_key") or "").strip()

    if not transcript or not notes:
        return jsonify({"error": "Both 'transcript' and 'notes' are required."}), 400

    # Sample transcript evenly: take beginning, middle, and end
    # This catches gaps across the entire lecture, not just the start
    t_words = transcript.split()
    t_total = len(t_words)

    if t_total <= 40_000:
        # Short enough — use the whole thing
        transcript_sample = transcript
        sample_desc       = f"full transcript ({t_total:,} words)"
    else:
        # Sample ~15k words from beginning, middle, and end
        # This gives coverage across the entire lecture
        sample_size = 15_000
        beginning   = " ".join(t_words[:sample_size])
        mid_start   = (t_total // 2) - (sample_size // 2)
        middle      = " ".join(t_words[mid_start:mid_start + sample_size])
        end         = " ".join(t_words[-sample_size:])
        transcript_sample = (
            f"[TRANSCRIPT BEGINNING]\n{beginning}\n\n"
            f"[TRANSCRIPT MIDDLE (around word {t_total//2:,})]\n{middle}\n\n"
            f"[TRANSCRIPT END]\n{end}"
        )
        sample_desc = f"sampled from {t_total:,} words (beginning + middle + end)"

    # Sample notes similarly if very large
    n_words = notes.split()
    n_total = len(n_words)
    if n_total <= 50_000:
        notes_sample = notes
        notes_desc   = f"full notes ({n_total:,} words)"
    else:
        # Beginning, middle, and end of notes
        ns = 16_000
        notes_sample = (
            f"[NOTES BEGINNING]\n{' '.join(n_words[:ns])}\n\n"
            f"[NOTES MIDDLE]\n{' '.join(n_words[n_total//2-ns//2:n_total//2+ns//2])}\n\n"
            f"[NOTES END]\n{' '.join(n_words[-ns:])}"
        )
        notes_desc = f"sampled from {n_total:,} words"

    system = (
        "You are a strict academic quality auditor. "
        "Your ONLY job is to verify that study notes completely cover the source lecture. "
        "Be mercilessly critical. If something is missing or thin, say so explicitly. "
        "Never give a verdict better than the actual coverage warrants. "
        "COMPREHENSIVE means 90%+ with no important gaps. Be strict about this threshold."
    )

    prompt = f"""Audit these study notes for completeness against the source lecture transcript.

TRANSCRIPT SAMPLE ({sample_desc}):
{transcript_sample}

NOTES SAMPLE ({notes_desc}):
{notes_sample}

═══════════════════════════════════
AUDIT REPORT — include ALL sections:
═══════════════════════════════════

## 1. Concept Extraction
From the TRANSCRIPT, identify and list 25 specific, distinct concepts, facts,
techniques, formulas, or examples. Be specific:
✗ BAD: "gradient descent was explained"
✓ GOOD: "the learning rate hyperparameter controls step size; too high causes divergence,
         too low causes slow convergence; typical values are 0.001 to 0.1"

## 2. Coverage Check
For each of the 25 items:
- ✅ COVERED — notes explain this clearly with sufficient depth
- ⚠️ PARTIAL — notes mention it but miss important details
- ❌ MISSING — notes do not cover this at all

## 3. Coverage Score
"X out of 25 concepts fully covered (Y%)"
Be strict: ⚠️ PARTIAL counts as 0.5, not 1.0

## 4. Thin Sections
List topics the instructor spent significant time on but notes only briefly mention.
Include specific page/section references from the notes.

## 5. Missing Topics
List important topics from the transcript completely absent from the notes.

## 6. Accuracy Check
Any statements in notes that contradict the transcript? List them explicitly.

## 7. Verdict
Choose ONE based on strict criteria:
- **COMPREHENSIVE** — >90% covered, no important gaps (suitable for exam prep alone)
- **GOOD** — 75-90% covered, minor gaps only (mostly suitable for self-study)
- **INCOMPLETE** — 50-75% covered, significant gaps (needs supplementation)
- **INADEQUATE** — <50% covered, major missing sections (needs full regeneration)

## 8. Recommended Actions
If verdict is not COMPREHENSIVE, give 3-5 specific, actionable steps to improve coverage.
Example: "Re-process Part 3 of the transcript (around words 45,000-60,000) which covers
backpropagation — this section appears thin in the current notes."
"""

    try:
        report = call_gemini(system, prompt, model_req, api_key)
        return jsonify({
            "report":                   report,
            "model_used":               model_req,
            "transcript_words_total":   t_total,
            "transcript_words_checked": word_count(transcript_sample),
            "notes_words_total":        n_total,
            "notes_words_checked":      word_count(notes_sample),
        })
    except Exception as e:
        logger.error("Verify error: %s", e)
        return jsonify({"error": str(e)}), 502


# ── Export routes ─────────────────────────────────────────────────────

@app.route("/api/export/pdf", methods=["POST"])
def api_export_pdf():
    """
    Convert HTML or Markdown content to a PDF file.

    Request body:
        {
            "content":  str,  — HTML or Markdown content
            "title":    str,  — document title
            "is_html":  bool  — True if content is HTML (default True)
        }

    Response: PDF file download
    """
    data       = request.get_json(silent=True) or {}
    content    = (data.get("content") or "").strip()
    title      = (data.get("title") or "LectureAI Notes").strip()
    is_html    = data.get("is_html", True)

    if not content:
        return jsonify({"error": "Missing 'content'."}), 400

    try:
        # If content is Markdown, convert to HTML first
        if not is_html:
            html_body = markdown_to_html_body(content)
        else:
            html_body = content

        pdf_bytes = generate_pdf_bytes(html_body, title)

        # Create safe filename from title
        safe_name = re.sub(r"[^a-zA-Z0-9\s\-_]", "", title)
        safe_name = re.sub(r"\s+", "-", safe_name).strip("-")[:60] or "notes"

        return send_file(
            io.BytesIO(pdf_bytes),
            mimetype         = "application/pdf",
            as_attachment    = True,
            download_name    = f"{safe_name}.pdf",
        )
    except RuntimeError as e:
        # WeasyPrint not installed or failed
        return jsonify({"error": str(e)}), 503
    except Exception as e:
        logger.exception("PDF export error")
        return jsonify({"error": f"PDF generation failed: {e}"}), 500


@app.route("/api/export/docx", methods=["POST"])
def api_export_docx():
    """
    Convert Markdown or HTML content to a Word DOCX file.

    Request body:
        {
            "content":     str,  — Markdown or HTML content
            "title":       str,  — document title
            "is_markdown": bool  — True if content is Markdown (default True)
        }

    Response: DOCX file download
    """
    data        = request.get_json(silent=True) or {}
    content     = (data.get("content") or "").strip()
    title       = (data.get("title") or "LectureAI Notes").strip()
    is_markdown = data.get("is_markdown", True)

    if not content:
        return jsonify({"error": "Missing 'content'."}), 400

    try:
        docx_bytes = generate_docx_bytes(content, title, is_markdown=is_markdown)

        safe_name = re.sub(r"[^a-zA-Z0-9\s\-_]", "", title)
        safe_name = re.sub(r"\s+", "-", safe_name).strip("-")[:60] or "notes"

        return send_file(
            io.BytesIO(docx_bytes),
            mimetype         = "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            as_attachment    = True,
            download_name    = f"{safe_name}.docx",
        )
    except RuntimeError as e:
        # python-docx not installed or failed
        return jsonify({"error": str(e)}), 503
    except Exception as e:
        logger.exception("DOCX export error")
        return jsonify({"error": f"DOCX generation failed: {e}"}), 500


def markdown_to_html_body(md: str) -> str:
    """
    Convert simple Markdown to HTML body content for PDF export.
    Handles headings, lists, code blocks, blockquotes, and inline formatting.
    """
    # Code blocks first (must come before inline code)
    def replace_code_block(m):
        lang = m.group(1) or ""
        code = html.escape(m.group(2).strip())
        return f'<pre><code class="language-{lang}">{code}</code></pre>'

    result = re.sub(r"```(\w*)\n?([\s\S]*?)```", replace_code_block, md)

    # Headings
    result = re.sub(r"^#### (.+)$", r"<h4>\1</h4>", result, flags=re.MULTILINE)
    result = re.sub(r"^### (.+)$",  r"<h3>\1</h3>", result, flags=re.MULTILINE)
    result = re.sub(r"^## (.+)$",   r"<h2>\1</h2>", result, flags=re.MULTILINE)
    result = re.sub(r"^# (.+)$",    r"<h1>\1</h1>", result, flags=re.MULTILINE)

    # Blockquotes
    result = re.sub(r"^> (.+)$", r"<blockquote>\1</blockquote>", result, flags=re.MULTILINE)

    # Lists
    result = re.sub(r"^[\*\-] (.+)$", r"<li>\1</li>", result, flags=re.MULTILINE)
    result = re.sub(r"^\d+\. (.+)$",  r"<li>\1</li>", result, flags=re.MULTILINE)

    # Inline formatting
    result = re.sub(r"\*\*(.+?)\*\*", r"<strong>\1</strong>", result)
    result = re.sub(r"\*(.+?)\*",     r"<em>\1</em>",         result)
    result = re.sub(r"`([^`]+)`",     r"<code>\1</code>",     result)

    # Wrap consecutive <li> tags in <ul>
    result = re.sub(r"(<li>[\s\S]*?</li>\n?)+",
                    lambda m: f"<ul>{m.group()}</ul>", result)

    # Wrap remaining text blocks in <p> tags
    lines_out = []
    for block in result.split("\n\n"):
        block = block.strip()
        if not block:
            continue
        if re.match(r"^<(h[1-6]|ul|ol|li|blockquote|pre)", block):
            lines_out.append(block)
        else:
            lines_out.append(f"<p>{block}</p>")

    return "\n".join(lines_out)


# ═══════════════════════════════════════════════════════════════════════
# SECTION 9: PLAYLIST SUPPORT
#
# Process an entire YouTube playlist — fetches every video's transcript,
# merges them all, then generates one comprehensive course notes document.
#
# How it works:
#   1. POST /api/playlist/info  → fetch playlist metadata and video list
#   2. POST /api/playlist/process/async → start background job that:
#        a. Fetches each video's transcript
#        b. Generates notes for each video
#        c. Merges all videos' notes into one course document
#   3. GET  /api/job/<job_id>   → poll for progress (reuses job system)
# ═══════════════════════════════════════════════════════════════════════

def extract_playlist_id(url: str):
    """
    Extract a YouTube playlist ID from a URL.

    Handles formats:
      - https://www.youtube.com/playlist?list=PLxxxxxx
      - https://www.youtube.com/watch?v=xxx&list=PLxxxxxx
      - PLxxxxxx (raw ID)
    """
    patterns = [
        r"[?&]list=([A-Za-z0-9_-]{10,})",  # list= parameter in any YouTube URL
        r"^(PL[A-Za-z0-9_-]{10,})$",        # raw playlist ID starting with PL
    ]
    for pattern in patterns:
        match = re.search(pattern, url.strip())
        if match:
            return match.group(1)
    return None


def fetch_playlist_info(playlist_id: str) -> dict:
    """
    Fetch metadata and video list for a YouTube playlist using yt-dlp.

    Returns:
        {
            "playlist_id":    str,
            "title":          str,   — playlist title
            "channel":        str,   — channel name
            "video_count":    int,   — total videos in playlist
            "videos": [
                {"id": str, "title": str, "duration_seconds": int},
                ...
            ]
        }

    Raises ValueError if yt-dlp fails or playlist not found.
    """
    url = f"https://www.youtube.com/playlist?list={playlist_id}"

    cmd = [
        "yt-dlp",
        "--flat-playlist",      # don't download, just list
        "--dump-single-json",   # output playlist info as JSON
        "--no-warnings",
        "--quiet",
        url,
    ]

    if COOKIES_FILE and Path(COOKIES_FILE).exists():
        cmd += ["--cookies", COOKIES_FILE]

    try:
        result = subprocess.run(
            cmd,
            capture_output=True,
            text=True,
            timeout=60,
        )
    except FileNotFoundError:
        raise ValueError("yt-dlp is not installed. Run: pip install yt-dlp")
    except subprocess.TimeoutExpired:
        raise ValueError("yt-dlp timed out fetching playlist info.")

    if result.returncode != 0:
        err = result.stderr.strip()
        if "does not exist" in err or "unavailable" in err.lower():
            raise ValueError("Playlist not found or is private.")
        raise ValueError(f"yt-dlp playlist error: {err[:200]}")

    if not result.stdout.strip():
        raise ValueError("yt-dlp returned no data for this playlist.")

    try:
        info = json.loads(result.stdout)
    except json.JSONDecodeError as e:
        raise ValueError(f"Could not parse playlist data: {e}")

    # Extract video list — yt-dlp returns 'entries' for playlists
    entries   = info.get("entries", [])
    video_list = []
    for entry in entries:
        if entry and entry.get("id"):
            video_list.append({
                "id":               entry.get("id", ""),
                "title":            entry.get("title", "Untitled"),
                "duration_seconds": entry.get("duration") or 0,
            })

    if not video_list:
        raise ValueError("Playlist appears to be empty or all videos are private.")

    return {
        "playlist_id":  playlist_id,
        "title":        info.get("title", "YouTube Playlist"),
        "channel":      info.get("uploader", ""),
        "video_count":  len(video_list),
        "videos":       video_list,
    }


def _run_playlist_job(
    job_id:      str,
    playlist_id: str,
    video_ids:   list,
    titles:      dict,
    course_title: str,
    model:       str,
    tone:        str,
    api_key:     str = "",
):
    """
    Background thread for playlist processing.
    Fetches transcript and generates notes for each video IN ORDER,
    then merges everything into one well-structured course document.
    
    Progress bar: 0-35% = fetching transcripts, 35-85% = generating notes, 85-98% = merging
    """
    try:
        _update_job(job_id, status="running", model_used=model)
        total_videos = len(video_ids)
        _add_progress(job_id, f"🎬 Starting playlist: {total_videos} videos | model={model}")
        _add_progress(job_id, f"📋 Phase 1/3: Fetching transcripts for all {total_videos} videos…")

        all_transcripts = []   # (video_index, title, transcript) — preserves order
        failed_videos   = []

        # ── Phase 1: Fetch all transcripts (0–35%) ────────────────
        for i, vid_id in enumerate(video_ids):
            vid_title = titles.get(vid_id, f"Video {i+1}")
            pct = int(((i) / total_videos) * 35)
            _update_job(job_id, percent=pct, chunks_done=i)
            _add_progress(job_id, f"[{i+1}/{total_videos}] Fetching: {vid_title[:55]}…")

            try:
                result     = fetch_transcript(vid_id)
                transcript = result["transcript"]
                wc         = result["word_count"]
                all_transcripts.append((i, vid_title, transcript))
                _add_progress(job_id, f"  ✓ {wc:,} words — OK")
            except Exception as e:
                failed_videos.append(vid_title)
                _add_progress(job_id, f"  ✗ SKIPPED (no captions): {str(e)[:60]}")
                logger.warning("[playlist %s] Video %s failed: %s", job_id[:8], vid_id, e)

        if not all_transcripts:
            raise ValueError(
                "Could not fetch any transcripts from this playlist. "
                "Make sure the videos have captions enabled."
            )

        _add_progress(job_id,
            f"✅ Phase 1 done: {len(all_transcripts)} fetched, {len(failed_videos)} skipped.")
        _update_job(job_id, percent=35)

        # ── Phase 2: Generate notes for each video (35–85%) ───────
        _add_progress(job_id, f"📝 Phase 2/3: Generating notes for {len(all_transcripts)} videos…")
        all_notes = []  # (original_index, title, notes) — preserves playlist order

        notes_range = 85 - 35  # 50% of bar for notes generation

        for seq_i, (orig_idx, vid_title, transcript) in enumerate(all_transcripts):
            pct = 35 + int((seq_i / len(all_transcripts)) * notes_range)
            _update_job(job_id, percent=pct, chunks_done=35 + seq_i)
            _add_progress(job_id,
                f"[{seq_i+1}/{len(all_transcripts)}] Notes: {vid_title[:55]}…")

            def on_progress(msg, _title=vid_title):
                _add_progress(job_id, f"    {msg}")

            try:
                result = process_transcript(
                    transcript  = transcript,
                    mode        = "notes",
                    tone_or_fmt = tone,
                    title       = vid_title,
                    model       = model,
                    api_key     = api_key,
                    progress_cb = on_progress,
                )
                all_notes.append((orig_idx, vid_title, result["content"]))
                _add_progress(job_id, f"  ✓ {result['words_out']:,} words of notes")
            except Exception as e:
                _add_progress(job_id, f"  ✗ Notes failed: {str(e)[:80]}")
                logger.warning("[playlist %s] Notes failed for %s: %s",
                               job_id[:8], vid_title, e)

        if not all_notes:
            raise ValueError("Failed to generate notes for any video in the playlist.")

        # Sort by original playlist order (important!)
        all_notes.sort(key=lambda x: x[0])

        _add_progress(job_id,
            f"✅ Phase 2 done: {len(all_notes)} videos have notes.")
        _update_job(job_id, percent=85)

        # ── Phase 3: Merge into ordered course document (85–98%) ──
        _add_progress(job_id,
            f"🔀 Phase 3/3: Merging {len(all_notes)} videos into one course document…")

        if len(all_notes) == 1:
            final_notes = all_notes[0][2]
        else:
            # Label each video's notes with its number and title, preserving ORDER
            labelled_notes = [
                f"# Video {i+1}: {title}\n\n{notes}"
                for i, (_, title, notes) in enumerate(all_notes)
            ]

            merge_system = (
                "You are an expert technical educator creating a comprehensive course document. "
                "Merge notes from multiple lecture videos into one complete, well-organized "
                "course textbook. Preserve the ORIGINAL VIDEO ORDER as the learning progression. "
                "Keep ALL content from every video's notes."
            )

            parts_text = "\n\n════════════════ VIDEO BREAK ════════════════\n\n".join(
                labelled_notes
            )
            course_merge_prompt = f"""Merge these notes from {len(all_notes)} lecture videos into ONE
comprehensive course document for: "{course_title}"

IMPORTANT: Preserve the VIDEO ORDER as the natural learning progression.

REQUIRED STRUCTURE (in this exact order):
1. ## Course Overview — 2-3 paragraphs summarizing the full course arc
2. ## Table of Contents — every major topic from every video, in video order
3. ## Prerequisites — what background knowledge is assumed
4. [Video 1 Topics] → [Video 2 Topics] → ... (keep video-by-video order for pedagogical flow)
   Use "## Part N: [Video Title]" headings to clearly mark each video's section
5. ## Complete Course Summary — comprehensive recap of everything covered
6. ## Master Reference Sheet — key formulas, algorithms, terms from the whole course

RULES:
- Keep EVERY concept, formula, example, code snippet, and analogy from every video
- Do NOT reorganize by topic — keep the video-by-video order (it's the intended curriculum)
- Each video section should be clearly labelled "## Part N: [Video Title]"
- Add brief transitions between video sections

INPUT NOTES ({len(all_notes)} videos in playlist order):
{parts_text[:150_000]}

Begin the course document:"""

            _update_job(job_id, percent=88)
            final_notes = call_gemini(merge_system, course_merge_prompt, model, api_key)

        _update_job(job_id, percent=98)

        # ── Build result ────────────────────────────────────────────
        skipped_note = (
            f"\n\n---\n*⚠ {len(failed_videos)} video(s) skipped (no captions): "
            f"{', '.join(failed_videos[:5])}{'…' if len(failed_videos) > 5 else ''}*"
            if failed_videos else ""
        )

        final_content = final_notes + skipped_note
        total_wc = sum(word_count(t) for _, _, t in all_transcripts)

        _update_job(
            job_id,
            status  = "done",
            percent = 100,
            result  = {
                "content":       final_content,
                "chunks":        len(all_notes),
                "words_in":      total_wc,
                "words_out":    word_count(final_content),
                "model_used":   model,
                "videos_done":  len(all_notes),
                "videos_failed": len(failed_videos),
            }
        )
        _add_progress(
            job_id,
            f"✓ Course notes complete: {len(all_notes)} videos, "
            f"{word_count(final_content):,} words output"
        )

    except Exception as exc:
        logger.exception("[playlist job %s] Failed", job_id[:8])
        _add_progress(job_id, f"ERROR: {exc}")
        _update_job(job_id, status="error", error=str(exc))


@app.route("/api/playlist/info", methods=["POST"])
def api_playlist_info():
    """
    Fetch metadata for a YouTube playlist without processing it.

    Used by the frontend to show the user what's in the playlist
    before they commit to processing all videos.

    Request body:
        { "url": "https://youtube.com/playlist?list=PLxxx" }

    Response:
        {
            "playlist_id": str,
            "title":       str,
            "channel":     str,
            "video_count": int,
            "videos": [{"id": str, "title": str, "duration_seconds": int}]
        }
    """
    data = request.get_json(silent=True) or {}
    url  = (data.get("url") or "").strip()

    if not url:
        return jsonify({"error": "Missing 'url' field."}), 400

    # Accept both playlist URLs and video URLs with list= parameter
    playlist_id = extract_playlist_id(url)
    if not playlist_id:
        return jsonify({"error": "No playlist ID found in URL. "
                                 "Make sure the URL contains '?list=PL...'"}), 400

    try:
        info = fetch_playlist_info(playlist_id)
        return jsonify(info)
    except ValueError as e:
        return jsonify({"error": str(e)}), 422
    except Exception as e:
        logger.exception("Playlist info error")
        return jsonify({"error": f"Server error: {e}"}), 500


@app.route("/api/playlist/process", methods=["POST"])
def api_playlist_process():
    """
    Start a background job to process a YouTube playlist.

    Fetches every video's transcript, generates notes for each,
    then merges everything into one comprehensive course notes document.

    Request body:
        {
            "playlist_id": str,      — playlist ID from /api/playlist/info
            "video_ids":   [str],    — list of video IDs to process (can be subset)
            "titles":      {id: str},— video ID → title mapping
            "course_title": str,     — playlist/course title
            "model":       str,      — Gemini model to use
            "tone":        str       — note style (e.g., "comprehensive")
        }

    Response: { "job_id": str, "video_count": int, "estimated_min": int }
    """
    data         = request.get_json(silent=True) or {}
    playlist_id  = (data.get("playlist_id") or "").strip()
    video_ids    = data.get("video_ids") or []
    titles       = data.get("titles") or {}
    course_title = (data.get("course_title") or "YouTube Course").strip()
    model_req    = (data.get("model") or "").strip() or GEMINI_MODEL
    tone         = (data.get("tone") or "comprehensive").strip()
    api_key      = (data.get("api_key") or "").strip()

    if not video_ids:
        return jsonify({"error": "No video IDs provided."}), 400
    # No hard limit — process the whole playlist!
    # Very large playlists (100+ videos) will just take longer.
    if len(video_ids) > 500:
        return jsonify({"error": "Maximum 500 videos per job."}), 400

    # Estimate: ~3-5 min/video (fetch + notes generation)
    is_flash      = "flash" in model_req.lower()
    min_per_video = 3 if is_flash else 5
    estimated_min = max(2, len(video_ids) * min_per_video)

    job_id = _create_job()
    _update_job(job_id, chunks_total=len(video_ids))  # "chunks" = videos for playlist

    thread = threading.Thread(
        target   = _run_playlist_job,
        args     = (job_id, playlist_id, video_ids, titles, course_title, model_req, tone, api_key),
        daemon   = True,
        name     = f"playlist-{job_id[:8]}"
    )
    thread.start()

    logger.info(
        "Playlist job started: id=%s videos=%d model=%s",
        job_id[:8], len(video_ids), model_req
    )

    return jsonify({
        "job_id":        job_id,
        "video_count":   len(video_ids),
        "estimated_min": estimated_min,
        "estimated_max": estimated_min * 2,
    })


# ═══════════════════════════════════════════════════════════════════════
# SECTION 10: APPLICATION ENTRY POINT
# ═══════════════════════════════════════════════════════════════════════

if __name__ == "__main__":
    # Warn if API key is missing (app will fail on first AI call)
    if not GEMINI_API_KEY:
        logger.warning("⚠  GEMINI_API_KEY not set — add it to your .env file.")

    # Log startup configuration
    logger.info(
        "Starting LectureAI | Model: %s | Chunk: %d words | "
        "Overlap: %d words | Async threshold: %d words | Max output: %d tokens",
        GEMINI_MODEL, CHUNK_WORDS, CHUNK_OVERLAP_WORDS,
        ASYNC_THRESHOLD_WORDS, MAX_OUTPUT_TOKENS
    )

    # Run the Flask development server
    # In production, use: gunicorn --config gunicorn.conf.py app:app
    app.run(host="0.0.0.0", port=8080, debug=False)
